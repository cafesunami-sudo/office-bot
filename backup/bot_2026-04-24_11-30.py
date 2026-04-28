from aiogram import Bot, Dispatcher
from aiogram.types import Message, ReplyKeyboardMarkup, KeyboardButton, FSInputFile
from aiogram.filters import Command
from docx import Document
from datetime import datetime, timedelta
import asyncio
import os
import json
import re

TOKEN = "8690185918:AAGOcbsIN_Kvom4dulvIk8XvQonQ7nvxIMs"

EMPLOYEES_FILE = r"D:\Проект офис БООТ\employees\sotrudniki.docx"
READY_FOLDER = r"D:\Проект офис БООТ\ready"
HISTORY_FILE = r"D:\Проект офис БООТ\history.json"

TEMPLATES = {
    "🌴 Полный отпуск": r"D:\Проект офис БООТ\templates\otpusk_full.docx",
    "🧩 Часть отпуска": r"D:\Проект офис БООТ\templates\otpusk_part.docx",
    "📌 Оставшийся отпуск": r"D:\Проект офис БООТ\templates\otpusk_rest.docx",
    "📝 БС с периода по период": r"D:\Проект офис БООТ\templates\bs_range.docx",
    "📅 БС на один день": r"D:\Проект офис БООТ\templates\bs_one.docx",
    "💍 Мат помощь (свадьба)": r"D:\Проект офис БООТ\templates\mat_wedding.docx",
    "👶 Мат помощь (ребенок)": r"D:\Проект офис БООТ\templates\mat_child.docx"
}

bot = Bot(token=TOKEN)
dp = Dispatcher()

data = {}
state = {}
temp_search = {}


def is_valid_date(text):
    if not re.fullmatch(r"\d{2}\.\d{2}\.\d{4}", text):
        return False
    try:
        datetime.strptime(text, "%d.%m.%Y")
        return True
    except:
        return False


def format_date_text(date_text):
    months = {
        "01": "января", "02": "февраля", "03": "марта", "04": "апреля",
        "05": "мая", "06": "июня", "07": "июля", "08": "августа",
        "09": "сентября", "10": "октября", "11": "ноября", "12": "декабря"
    }
    dt = datetime.strptime(date_text, "%d.%m.%Y")
    return f"{dt.day} {months[dt.strftime('%m')]} {dt.year}"


def format_fio_short(fio):
    parts = fio.split()
    if len(parts) >= 3:
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    return fio


def load_employees():
    doc = Document(EMPLOYEES_FILE)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def load_history():
    if not os.path.exists(HISTORY_FILE):
        return []
    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except:
        return []


def save_history(d):
    record = {
        "fio": d.get("fio"),
        "position": d.get("pos"),
        "type": d.get("type"),
        "start": d.get("start", ""),
        "end": d.get("end", ""),
        "days": d.get("days", ""),
        "created_at": datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    }

    history = load_history()
    history.append(record)

    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=2)


def replace_text(doc, rep):
    for p in doc.paragraphs:
        text = p.text
        for k, v in rep.items():
            text = text.replace(k, v)
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    for k, v in rep.items():
                        text = text.replace(k, v)
                    if p.runs:
                        p.runs[0].text = text
                        for r in p.runs[1:]:
                            r.text = ""


def create_doc(d):
    doc = Document(TEMPLATES[d["type"]])

    rep = {
        "{{FIO}}": format_fio_short(d["fio"]),
        "{{POSITION}}": d["pos"],
        "{{TODAY}}": format_date_text(datetime.now().strftime("%d.%m.%Y"))
    }

    if "start" in d:
        rep["{{DATE_START}}"] = format_date_text(d["start"])

    if "end" in d:
        rep["{{DATE_END}}"] = format_date_text(d["end"])

    replace_text(doc, rep)

    os.makedirs(READY_FOLDER, exist_ok=True)
    path = os.path.join(READY_FOLDER, f"{d['fio']}.docx")
    doc.save(path)
    return path


def finish_and_send(chat_id):
    path = create_doc(data[chat_id])
    save_history(data[chat_id])
    return path


def employee_keyboard(employees):
    kb = [[KeyboardButton(text=e)] for e in employees]
    kb.append([KeyboardButton(text="🏠 Старт")])
    return ReplyKeyboardMarkup(keyboard=kb, resize_keyboard=True)


menu = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text="📄 Создать заявление")],
    [KeyboardButton(text="📜 История")],
    [KeyboardButton(text="🏠 Старт")]
], resize_keyboard=True)

history_menu = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text="🔍 Поиск сотрудника")],
    [KeyboardButton(text="📋 Полный список сотрудников")],
    [KeyboardButton(text="🏠 Старт")]
], resize_keyboard=True)

pos_menu = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text="Инженер программист")],
    [KeyboardButton(text="Программист")],
    [KeyboardButton(text="🏠 Старт")]
], resize_keyboard=True)


@dp.message(Command("start"))
async def start(m: Message):
    state[m.chat.id] = "menu"
    await m.answer("Меню", reply_markup=menu)


@dp.message()
async def handler(m: Message):
    chat_id = m.chat.id
    text = m.text

    if text == "🏠 Старт":
        state[chat_id] = "menu"
        await m.answer("Меню", reply_markup=menu)
        return

    if text == "📄 Создать заявление":
        state[chat_id] = "search_emp"
        await m.answer("Напиши ФИО или часть ФИО сотрудника")
        return

    if state.get(chat_id) == "search_emp":
        employees = load_employees()
        found = [e for e in employees if text.lower() in e.lower()]

        if not found:
            await m.answer("Сотрудник не найден. Напиши другую часть ФИО.")
            return

        temp_search[chat_id] = found
        state[chat_id] = "choose_emp"
        await m.answer("Выбери сотрудника:", reply_markup=employee_keyboard(found))
        return

    if state.get(chat_id) == "choose_emp":
        if text not in temp_search.get(chat_id, []):
            await m.answer("Выбери сотрудника только из списка кнопок.")
            return

        data[chat_id] = {"fio": text}
        state[chat_id] = "type"

        kb = [[KeyboardButton(text=t)] for t in TEMPLATES]
        kb.append([KeyboardButton(text="🏠 Старт")])
        await m.answer("Выбери тип заявления:", reply_markup=ReplyKeyboardMarkup(keyboard=kb, resize_keyboard=True))
        return

    if state.get(chat_id) == "type":
        if text not in TEMPLATES:
            await m.answer("Выбери тип заявления из списка.")
            return

        data[chat_id]["type"] = text
        state[chat_id] = "pos"
        await m.answer("Должность:", reply_markup=pos_menu)
        return

    if state.get(chat_id) == "pos":
        if text not in ["Инженер программист", "Программист"]:
            await m.answer("Выбери должность из списка.")
            return

        data[chat_id]["pos"] = text

        if data[chat_id]["type"] == "👶 Мат помощь (ребенок)":
            path = finish_and_send(chat_id)
            await m.answer_document(FSInputFile(path))
            state[chat_id] = "menu"
            await m.answer("Готово", reply_markup=menu)
            return

        state[chat_id] = "date"
        await m.answer("Введи дату начала ДД.ММ.ГГГГ")
        return

    if state.get(chat_id) == "date":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату в формате ДД.ММ.ГГГГ")
            return

        data[chat_id]["start"] = text
        start_date = datetime.strptime(text, "%d.%m.%Y")
        doc_type = data[chat_id]["type"]

        if doc_type == "🌴 Полный отпуск":
            days = 30
            end_date = start_date + timedelta(days=days - 1)
            data[chat_id]["days"] = str(days)
            data[chat_id]["end"] = end_date.strftime("%d.%m.%Y")

            path = finish_and_send(chat_id)
            await m.answer_document(FSInputFile(path))
            state[chat_id] = "menu"
            await m.answer("Готово", reply_markup=menu)
            return

        if doc_type == "🧩 Часть отпуска":
            days = 15
            end_date = start_date + timedelta(days=days - 1)
            data[chat_id]["days"] = str(days)
            data[chat_id]["end"] = end_date.strftime("%d.%m.%Y")

            path = finish_and_send(chat_id)
            await m.answer_document(FSInputFile(path))
            state[chat_id] = "menu"
            await m.answer("Готово", reply_markup=menu)
            return

        if doc_type == "📅 БС на один день":
            data[chat_id]["days"] = "1"
            data[chat_id]["end"] = text

            path = finish_and_send(chat_id)
            await m.answer_document(FSInputFile(path))
            state[chat_id] = "menu"
            await m.answer("Готово", reply_markup=menu)
            return

        if doc_type == "💍 Мат помощь (свадьба)":
            days = 3
            end_date = start_date + timedelta(days=days - 1)
            data[chat_id]["days"] = str(days)
            data[chat_id]["end"] = end_date.strftime("%d.%m.%Y")

            path = finish_and_send(chat_id)
            await m.answer_document(FSInputFile(path))
            state[chat_id] = "menu"
            await m.answer("Готово", reply_markup=menu)
            return

        state[chat_id] = "days"
        await m.answer("Количество дней")
        return

    if state.get(chat_id) == "days":
        if not text.isdigit():
            await m.answer("Введи количество дней цифрой.")
            return

        days = int(text)
        start_date = datetime.strptime(data[chat_id]["start"], "%d.%m.%Y")
        end_date = start_date + timedelta(days=days - 1)

        data[chat_id]["days"] = str(days)
        data[chat_id]["end"] = end_date.strftime("%d.%m.%Y")

        path = finish_and_send(chat_id)
        await m.answer_document(FSInputFile(path))
        state[chat_id] = "menu"
        await m.answer("Готово", reply_markup=menu)
        return

    if text == "📜 История":
        state[chat_id] = "history_menu"
        await m.answer("Выбери действие:", reply_markup=history_menu)
        return

    if text == "🔍 Поиск сотрудника":
        state[chat_id] = "history_search"
        await m.answer("Напиши ФИО или часть ФИО сотрудника")
        return

    if text == "📋 Полный список сотрудников":
        employees = load_employees()
        temp_search[chat_id] = employees
        state[chat_id] = "history_choose"
        await m.answer("Выбери сотрудника:", reply_markup=employee_keyboard(employees))
        return

    if state.get(chat_id) == "history_search":
        employees = load_employees()
        found = [e for e in employees if text.lower() in e.lower()]

        if not found:
            await m.answer("Сотрудник не найден. Попробуй ещё раз.")
            return

        temp_search[chat_id] = found
        state[chat_id] = "history_choose"
        await m.answer("Выбери сотрудника:", reply_markup=employee_keyboard(found))
        return

    if state.get(chat_id) == "history_choose":
        if text not in temp_search.get(chat_id, []):
            await m.answer("Выбери сотрудника только из списка кнопок.")
            return

        history = load_history()
        records = [r for r in history if r.get("fio") == text]

        if not records:
            await m.answer("По этому сотруднику история пустая.", reply_markup=menu)
            state[chat_id] = "menu"
            return

        msg = ""
        for r in records:
            msg += f"ФИО: {r.get('fio')}\n"
            msg += f"Тип: {r.get('type')}\n"
            msg += f"Должность: {r.get('position')}\n"

            if r.get("start"):
                msg += f"Начало: {r.get('start')}\n"

            if r.get("end"):
                msg += f"Конец: {r.get('end')}\n"

            if r.get("days"):
                msg += f"Дней: {r.get('days')}\n"

            msg += f"Создан: {r.get('created_at')}\n\n"

        await m.answer(msg, reply_markup=menu)
        state[chat_id] = "menu"
        return


async def main():
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())