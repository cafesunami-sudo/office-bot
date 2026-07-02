from aiogram import Bot, Dispatcher
from aiogram.types import Message, ReplyKeyboardMarkup, KeyboardButton, FSInputFile, Update
from aiogram.filters import Command
from aiohttp import web
from docx import Document
from datetime import datetime, timedelta
from calendar import monthrange
from zoneinfo import ZoneInfo
import asyncio
import os
import json
import re
import uuid

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment
except Exception:
    Workbook = None
    Font = None
    Alignment = None

try:
    import psycopg2
    from psycopg2.extras import Json
except Exception:
    psycopg2 = None
    Json = None

TIMEZONE = ZoneInfo(os.environ.get("TIMEZONE", "Asia/Tashkent"))


def now_dt():
    return datetime.now(TIMEZONE)


TOKEN = os.environ.get("BOT_TOKEN")

ADMIN_CHAT_ID = int(os.environ.get("ADMIN_CHAT_ID", 0))
GROUP_CHAT_ID = int(os.environ.get("GROUP_CHAT_ID", 0))

ADMIN_USERS = [137602775]
REPORT_ONLY_USERS = [390687401]

# ВАЖНО: оставил как ты просил ранее
REMIND_TIMES = ["09:30"]
START_REMIND_TIME = "17:00"

EMPLOYEES_FILE = "employees/sotrudniki.docx"
SALARY_FILE = "employees/salary.docx"
READY_FOLDER = "ready"
HISTORY_FILE = "history.json"
HOLIDAYS_FILE = "holidays.json"
REMINDERS_SENT_FILE = "reminders_sent.json"
DATABASE_URL = os.environ.get("DATABASE_URL")

TEMPLATES = {
    "🌴 Полный отпуск": "templates/otpusk_full.docx",
    "🧩 Часть отпуска": "templates/otpusk_part.docx",
    "📌 Оставшийся отпуск": "templates/otpusk_rest.docx",
    "📚 Учебный отпуск": "templates/study_leave.docx",
    "📝 БС с периода по период": "templates/bs_range.docx",
    "📅 БС на один день": "templates/bs_one.docx",
    "💍 Мат помощь (свадьба)": "templates/mat_wedding.docx",
    "👶 Мат помощь (ребенок)": "templates/mat_child.docx",
    "👶 Мат помощь (ребенок + 3 дня)": "templates/mat_child_3_days.docx",
    "🕊 Мат помощь (смерть родственника)": "templates/mat_death.docx",
}

TRIP_REPORT_TEMPLATE = "templates/trip/business_trip_report.docx"
TRIP_CERTIFICATE_TEMPLATE = "templates/trip/business_trip_certificate.docx"
TRIP_TYPE = "✈️ Командировка"
TRIP_REGION_SUFFIX = "подразделения по Хорезмской области"

SICK_LEAVE_TYPE = "🏥 Больничный"
BS_LEAVE_TYPES = ["📝 БС с периода по период", "📅 БС на один день"]
BS_RANGE_TYPE = "📝 БС с периода по период"
CHILD_BIRTH_3_DAYS_TYPE = "👶 Мат помощь (ребенок + 3 дня)"

# Эти заявления нужны только для формирования документа и истории.
# По ним бот НЕ должен отправлять сообщения в группу,
# НЕ должен отправлять напоминания о начале/выходе на работу
# и НЕ должен давать менять дату выхода через историю.
MATERIAL_ASSISTANCE_TYPES = [
    "💍 Мат помощь (свадьба)",
    "👶 Мат помощь (ребенок)",
    "🕊 Мат помощь (смерть родственника)",
]

# Эти виды матпомощи не требуют ввода даты периода: в документе ставится сегодняшняя дата.
MATERIAL_ASSISTANCE_NO_DATE_TYPES = [
    "👶 Мат помощь (ребенок)",
    "🕊 Мат помощь (смерть родственника)",
]

MANUAL_TYPES = [
    "🌴 Полный отпуск",
    "🧩 Часть отпуска",
    "📌 Оставшийся отпуск",
    "📚 Учебный отпуск",
    "📝 БС с периода по период",
    "📅 БС на один день",
    "🏥 Больничный",
]

bot = Bot(token=TOKEN)
dp = Dispatcher()

data = {}
state = {}
temp_search = {}

# ================== КЭШ ДЛЯ СКОРОСТИ ==================
# После ECO-режима Neon может "засыпать". Если каждый шаг формы читает Neon
# и заново открывает salary.docx, бот отвечает с задержкой 20-40 секунд.
# Поэтому историю, напоминания и salary.docx держим в памяти процесса.
HISTORY_CACHE = None
SENT_REMINDERS_CACHE = None
SALARY_RECORDS_CACHE = {
    "mtime": None,
    "records": None,
}

def reset_history_cache():
    global HISTORY_CACHE
    HISTORY_CACHE = None

def reset_sent_reminders_cache():
    global SENT_REMINDERS_CACHE
    SENT_REMINDERS_CACHE = None

def reset_salary_cache():
    global SALARY_RECORDS_CACHE
    SALARY_RECORDS_CACHE = {"mtime": None, "records": None}



def make_keyboard(buttons, cols=2):
    keyboard = []
    row = []
    for btn in buttons:
        row.append(KeyboardButton(text=btn))
        if len(row) == cols:
            keyboard.append(row)
            row = []
    if row:
        keyboard.append(row)
    return ReplyKeyboardMarkup(keyboard=keyboard, resize_keyboard=True)


def is_allowed(chat_id):
    return chat_id in ADMIN_USERS or chat_id in REPORT_ONLY_USERS


def is_admin(chat_id):
    return chat_id in ADMIN_USERS


def is_report_only(chat_id):
    return chat_id in REPORT_ONLY_USERS


def get_menu(chat_id):
    if is_admin(chat_id):
        return menu
    if is_report_only(chat_id):
        return report_only_menu
    return None


def normalize_date(text):
    text = str(text or "").strip()
    match = re.fullmatch(r"(\d{1,2})\.(\d{1,2})\.(\d{4})", text)
    if not match:
        return text
    day, month, year = match.groups()
    try:
        dt = datetime(int(year), int(month), int(day))
        return dt.strftime("%d.%m.%Y")
    except Exception:
        return text


def is_valid_date(text):
    text = normalize_date(text)
    if not re.fullmatch(r"\d{2}\.\d{2}\.\d{4}", text):
        return False
    try:
        datetime.strptime(text, "%d.%m.%Y")
        return True
    except Exception:
        return False


def parse_date_or_none(text):
    try:
        return datetime.strptime(normalize_date(text), "%d.%m.%Y").date()
    except Exception:
        return None


def load_holidays():
    if not os.path.exists(HOLIDAYS_FILE):
        return set()
    try:
        with open(HOLIDAYS_FILE, "r", encoding="utf-8") as f:
            return set(json.load(f))
    except Exception:
        return set()


HOLIDAYS = load_holidays()


def get_return_to_work_date(end_date_text):
    d = datetime.strptime(normalize_date(end_date_text), "%d.%m.%Y") + timedelta(days=1)
    while d.weekday() >= 5 or d.strftime("%Y-%m-%d") in HOLIDAYS:
        d += timedelta(days=1)
    return d.strftime("%d.%m.%Y")


def load_sent_reminders_from_json():
    if not os.path.exists(REMINDERS_SENT_FILE):
        return []
    try:
        with open(REMINDERS_SENT_FILE, "r", encoding="utf-8") as f:
            content = json.load(f)
            return content if isinstance(content, list) else []
    except Exception:
        return []


def save_sent_reminders_to_json(sent):
    with open(REMINDERS_SENT_FILE, "w", encoding="utf-8") as f:
        json.dump(sent, f, ensure_ascii=False, indent=2)



def load_sent_reminders():
    global SENT_REMINDERS_CACHE

    if SENT_REMINDERS_CACHE is not None:
        return list(SENT_REMINDERS_CACHE)

    if not db_enabled():
        SENT_REMINDERS_CACHE = load_sent_reminders_from_json()
        return list(SENT_REMINDERS_CACHE)

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT key FROM sent_reminders")
        rows = cur.fetchall()
        cur.close()
        SENT_REMINDERS_CACHE = [row[0] for row in rows]
        return list(SENT_REMINDERS_CACHE)
    except Exception as e:
        print("POSTGRES LOAD SENT REMINDERS ERROR:", e)
        SENT_REMINDERS_CACHE = load_sent_reminders_from_json()
        return list(SENT_REMINDERS_CACHE)
    finally:
        if conn:
            conn.close()


def save_sent_reminder(key):
    global SENT_REMINDERS_CACHE

    key = str(key or "").strip()
    if not key:
        return

    if SENT_REMINDERS_CACHE is not None and key not in SENT_REMINDERS_CACHE:
        SENT_REMINDERS_CACHE.append(key)

    if not db_enabled():
        sent = load_sent_reminders_from_json()
        if key not in sent:
            sent.append(key)
            save_sent_reminders_to_json(sent)
        SENT_REMINDERS_CACHE = sent
        return

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("INSERT INTO sent_reminders (key) VALUES (%s) ON CONFLICT (key) DO NOTHING", (key,))
        conn.commit()
        cur.close()
    except Exception as e:
        print("POSTGRES SAVE SENT REMINDER ERROR:", e)
        sent = load_sent_reminders_from_json()
        if key not in sent:
            sent.append(key)
            save_sent_reminders_to_json(sent)
        SENT_REMINDERS_CACHE = sent
    finally:
        if conn:
            conn.close()


def reserve_sent_reminder(key):
    global SENT_REMINDERS_CACHE

    """
    Ставит ключ отправки ДО отправки сообщения.
    Это защищает от дублей после deploy/restart:
    если ключ уже есть в PostgreSQL/JSON, сообщение повторно не отправляем.
    """
    key = str(key or "").strip()
    if not key:
        return False

    if not db_enabled():
        sent = load_sent_reminders_from_json()
        if key in sent:
            SENT_REMINDERS_CACHE = sent
            return False
        sent.append(key)
        save_sent_reminders_to_json(sent)
        SENT_REMINDERS_CACHE = sent
        return True

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO sent_reminders (key) VALUES (%s) ON CONFLICT (key) DO NOTHING RETURNING key",
            (key,)
        )
        inserted = cur.fetchone()
        conn.commit()
        cur.close()

        if inserted is not None:
            if SENT_REMINDERS_CACHE is None:
                SENT_REMINDERS_CACHE = []
            if key not in SENT_REMINDERS_CACHE:
                SENT_REMINDERS_CACHE.append(key)
            return True
        return False
    except Exception as e:
        print("POSTGRES RESERVE SENT REMINDER ERROR:", e)
        sent = load_sent_reminders_from_json()
        if key in sent:
            SENT_REMINDERS_CACHE = sent
            return False
        sent.append(key)
        save_sent_reminders_to_json(sent)
        SENT_REMINDERS_CACHE = sent
        return True
    finally:
        if conn:
            conn.close()

def migrate_json_reminders_to_postgres():
    if not db_enabled():
        return

    json_sent = load_sent_reminders_from_json()
    if not json_sent:
        print("PostgreSQL: reminders_sent.json пустой, переносить нечего")
        return

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        migrated = 0
        for key in json_sent:
            key = str(key or "").strip()
            if not key:
                continue
            cur.execute("INSERT INTO sent_reminders (key) VALUES (%s) ON CONFLICT (key) DO NOTHING", (key,))
            migrated += 1
        conn.commit()
        cur.close()
        print(f"PostgreSQL: перенесено ключей reminders_sent: {migrated}")
    except Exception as e:
        print("POSTGRES REMINDERS MIGRATION ERROR:", e)
    finally:
        if conn:
            conn.close()

def month_name(date_text):
    months = {
        "01": "января", "02": "февраля", "03": "марта", "04": "апреля",
        "05": "мая", "06": "июня", "07": "июля", "08": "августа",
        "09": "сентября", "10": "октября", "11": "ноября", "12": "декабря",
    }
    dt = datetime.strptime(normalize_date(date_text), "%d.%m.%Y")
    return months[dt.strftime("%m")]


def format_date_text(date_text):
    dt = datetime.strptime(normalize_date(date_text), "%d.%m.%Y")
    return f"«{dt.strftime('%d')}» {month_name(date_text)} {dt.year} года"


def format_date_text_without_year_word(date_text):
    dt = datetime.strptime(normalize_date(date_text), "%d.%m.%Y")
    return f"«{dt.strftime('%d')}» {month_name(date_text)} {dt.year}"


def format_date_range_start(date_text):
    dt = datetime.strptime(normalize_date(date_text), "%d.%m.%Y")
    return f"«{dt.strftime('%d')}» {month_name(date_text)} {dt.year}"


def format_date_range_end(date_text):
    dt = datetime.strptime(normalize_date(date_text), "%d.%m.%Y")
    return f"«{dt.strftime('%d')}» {month_name(date_text)} {dt.year} года"


def format_fio_short(fio):
    parts = str(fio or "").split()
    if len(parts) >= 3:
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    return fio


def type_uz(t):
    if t == "📌 Оставшийся отпуск":
        return "Mehnat ta’tilining qolgan qismi"
    if t in ["🌴 Полный отпуск", "🧩 Часть отпуска"]:
        return "Mehnat ta’tili"
    if t in ["📝 БС с периода по период", "📅 БС на один день"]:
        return "Ish haqi saqlanmagan ta’til"
    if t == "📚 Учебный отпуск":
        return "O‘quv ta’tili"
    if t == "🏥 Больничный":
        return "Kasallik ta’tili"
    if t == CHILD_BIRTH_3_DAYS_TYPE:
        return "Moddiy yordam va 3 kunlik ta’til"
    if t in MATERIAL_ASSISTANCE_TYPES:
        return "Moddiy yordam"
    return t


def start_phrase_uz(t):
    if t == "📌 Оставшийся отпуск":
        return "mehnat ta’tilining qolgan qismiga chiqadi"
    if t in ["🌴 Полный отпуск", "🧩 Часть отпуска"]:
        return "ta’tilga chiqadi"
    if t in ["📝 БС с периода по период", "📅 БС на один день"]:
        return "ish haqi saqlanmagan ta’tilga chiqadi"
    if t == "📚 Учебный отпуск":
        return "o‘quv ta’tiliga chiqadi"
    if t == "🏥 Больничный":
        return "kasallik ta’tiliga chiqadi"
    if t == CHILD_BIRTH_3_DAYS_TYPE:
        return "3 kunlik ta’tilga chiqadi"
    return "ta’tilga chiqadi"


def is_material_assistance(record_or_type):
    """True для заявлений на матпомощь: свадьба/рождение ребенка/смерть родственника."""
    if isinstance(record_or_type, dict):
        record_type = record_or_type.get("type", "")
    else:
        record_type = record_or_type
    return record_type in MATERIAL_ASSISTANCE_TYPES



def to_latin_uz(text):
    text = str(text or "")

    # Avval ko‘p harfli birikmalar
    replacements = [
        ("Ў", "O‘"), ("ў", "o‘"),
        ("Ғ", "G‘"), ("ғ", "g‘"),
        ("Қ", "Q"), ("қ", "q"),
        ("Ҳ", "H"), ("ҳ", "h"),
        ("Ё", "Yo"), ("ё", "yo"),
        ("Ю", "Yu"), ("ю", "yu"),
        ("Я", "Ya"), ("я", "ya"),
        ("Е", "Ye"), ("е", "e"),
        ("Ц", "Ts"), ("ц", "ts"),
        ("Ч", "Ch"), ("ч", "ch"),
        ("Ш", "Sh"), ("ш", "sh"),
        ("Щ", "Sh"), ("щ", "sh"),
        ("Ъ", ""), ("ъ", ""),
        ("Ь", ""), ("ь", ""),
    ]

    for src, dst in replacements:
        text = text.replace(src, dst)

    table = str.maketrans({
        "А": "A", "а": "a",
        "Б": "B", "б": "b",
        "В": "V", "в": "v",
        "Г": "G", "г": "g",
        "Д": "D", "д": "d",
        "Ж": "J", "ж": "j",
        "З": "Z", "з": "z",
        "И": "I", "и": "i",
        "Й": "Y", "й": "y",
        "К": "K", "к": "k",
        "Л": "L", "л": "l",
        "М": "M", "м": "m",
        "Н": "N", "н": "n",
        "О": "O", "о": "o",
        "П": "P", "п": "p",
        "Р": "R", "р": "r",
        "С": "S", "с": "s",
        "Т": "T", "т": "t",
        "У": "U", "у": "u",
        "Ф": "F", "ф": "f",
        "Х": "X", "х": "x",
        "Ы": "I", "ы": "i",
        "Э": "E", "э": "e",
    })

    return text.translate(table)


def group_value(text):
    return to_latin_uz(text)


def position_group_value(position):
    pos = str(position or "").strip()
    normalized = pos.lower().replace("-", " ")
    normalized = " ".join(normalized.split())

    if normalized == "программист":
        return "Dasturchi"
    if normalized == "инженер программист":
        return "Muhandis-dasturchi"

    return group_value(pos)


def load_employees():
    # Теперь основной список сотрудников берется из employees/salary.docx.
    # Старый файл employees/sotrudniki.docx больше не нужен для поиска ФИО.
    records = load_salary_records()
    if records:
        return [r.get("fio", "") for r in records if r.get("fio")]

    # Резервный вариант: если salary.docx временно отсутствует, бот не падает.
    try:
        doc = Document(EMPLOYEES_FILE)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except Exception:
        return []


def search_employees_by_text(text, employees):
    q = str(text or "").lower().strip()
    if len(q) < 2:
        return None
    return [e for e in employees if any(part.lower().startswith(q) for part in e.split())]


def fio_startswith_text(fio, text):
    q = str(text or "").lower().strip()
    if len(q) < 2:
        return False
    return any(part.lower().startswith(q) for part in str(fio or "").split())


def unique_list(values):
    result = []
    seen = set()
    for value in values:
        value = str(value or "").strip()
        if not value or value in seen:
            continue
        seen.add(value)
        result.append(value)
    return result


def db_enabled():
    return bool(DATABASE_URL and psycopg2 is not None)


def get_db_conn():
    if not db_enabled():
        return None
    return psycopg2.connect(DATABASE_URL)


def init_database():
    if not DATABASE_URL:
        print("DATABASE_URL не указан. История и напоминания будут храниться в JSON-файлах")
        return

    if psycopg2 is None:
        print("psycopg2 не установлен. Добавь psycopg2-binary в requirements.txt")
        return

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()

        cur.execute("""
            CREATE TABLE IF NOT EXISTS history_records (
                id SERIAL PRIMARY KEY,
                data JSONB NOT NULL,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS sent_reminders (
                key TEXT PRIMARY KEY,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)

        conn.commit()
        cur.close()
        print("PostgreSQL готов: history_records и sent_reminders проверены")
    except Exception as e:
        print("POSTGRES INIT ERROR:", e)
    finally:
        if conn:
            conn.close()

def load_history_from_json():
    if not os.path.exists(HISTORY_FILE):
        return []
    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            content = json.load(f)
            return content if isinstance(content, list) else []
    except Exception:
        return []


def save_history_to_json(history):
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=2)



def load_history():
    global HISTORY_CACHE

    if HISTORY_CACHE is not None:
        return [dict(r) for r in HISTORY_CACHE if isinstance(r, dict)]

    if not db_enabled():
        HISTORY_CACHE = load_history_from_json()
        return [dict(r) for r in HISTORY_CACHE if isinstance(r, dict)]

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT data FROM history_records ORDER BY id ASC")
        rows = cur.fetchall()
        cur.close()
        HISTORY_CACHE = [row[0] for row in rows if isinstance(row[0], dict)]
        return [dict(r) for r in HISTORY_CACHE]
    except Exception as e:
        print("POSTGRES LOAD HISTORY ERROR:", e)
        HISTORY_CACHE = load_history_from_json()
        return [dict(r) for r in HISTORY_CACHE if isinstance(r, dict)]
    finally:
        if conn:
            conn.close()


def save_history_full(history):
    global HISTORY_CACHE

    clean_history = [dict(r) for r in history if isinstance(r, dict)]
    HISTORY_CACHE = clean_history

    if not db_enabled():
        save_history_to_json(clean_history)
        return

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("DELETE FROM history_records")
        for record in clean_history:
            cur.execute("INSERT INTO history_records (data) VALUES (%s)", (Json(record),))
        conn.commit()
        cur.close()
    except Exception as e:
        print("POSTGRES SAVE HISTORY FULL ERROR:", e)
        save_history_to_json(clean_history)
    finally:
        if conn:
            conn.close()


def append_history_record(record):
    global HISTORY_CACHE

    if not isinstance(record, dict):
        return

    if HISTORY_CACHE is None:
        HISTORY_CACHE = load_history()
    HISTORY_CACHE.append(dict(record))

    if not db_enabled():
        history = load_history_from_json()
        history.append(record)
        save_history_to_json(history)
        return

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("INSERT INTO history_records (data) VALUES (%s)", (Json(record),))
        conn.commit()
        cur.close()
    except Exception as e:
        print("POSTGRES APPEND HISTORY ERROR:", e)
        history = load_history_from_json()
        history.append(record)
        save_history_to_json(history)
    finally:
        if conn:
            conn.close()

def migrate_json_history_to_postgres():
    if not db_enabled():
        return

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM history_records")
        count = cur.fetchone()[0]
        cur.close()
        conn.close()
        conn = None

        if count > 0:
            print("PostgreSQL: история уже есть, миграция не нужна")
            return

        json_history = load_history_from_json()
        if not json_history:
            print("PostgreSQL: history.json пустой, переносить нечего")
            return

        save_history_full(json_history)
        print(f"PostgreSQL: перенесено записей из history.json: {len(json_history)}")
    except Exception as e:
        print("POSTGRES MIGRATION ERROR:", e)
    finally:
        if conn:
            conn.close()


def get_history_fios():
    return unique_list([r.get("fio", "") for r in load_history() if isinstance(r, dict)])


def get_employees_and_history_fios():
    try:
        employees = load_employees()
    except Exception:
        employees = []
    return unique_list(employees + get_history_fios())


def search_fios_by_text(text, fios):
    q = str(text or "").lower().strip()
    if len(q) < 2:
        return None
    return [fio for fio in fios if any(part.lower().startswith(q) for part in str(fio).split())]


def normalize_created_at(value):
    if not value:
        return now_dt().strftime("%d.%m.%Y %H:%M:%S")
    value = str(value).strip()
    for fmt in ["%d.%m.%Y %H:%M:%S", "%d.%m.%Y %H:%M", "%d.%m.%Y"]:
        try:
            return datetime.strptime(value, fmt).strftime("%d.%m.%Y %H:%M:%S")
        except Exception:
            pass
    return now_dt().strftime("%d.%m.%Y %H:%M:%S")


def days_between(start_text, end_text):
    start_date = datetime.strptime(normalize_date(start_text), "%d.%m.%Y")
    end_date = datetime.strptime(normalize_date(end_text), "%d.%m.%Y")
    return (end_date - start_date).days + 1


def calc_total_days(periods):
    total = 0
    for p in periods:
        try:
            total += days_between(p.get("start", ""), p.get("end", ""))
        except Exception:
            pass
    return total


def get_periods_from_record(record):
    periods = record.get("periods")
    if isinstance(periods, list) and periods:
        return periods
    if record.get("start") and record.get("end"):
        return [{"start": record.get("start"), "end": record.get("end")}]
    return []


def format_periods(periods):
    return "\n".join([f"{i}. {p.get('start')} — {p.get('end')}" for i, p in enumerate(periods, 1)])


def normalize_history_record(record):
    fixed = dict(record)
    fixed["fio"] = fixed.get("fio", "")
    fixed["position"] = fixed.get("position", "")
    fixed["project"] = fixed.get("project", "")
    fixed["type"] = fixed.get("type", "")
    fixed["start"] = normalize_date(fixed.get("start", "")) if fixed.get("start") else ""
    fixed["end"] = normalize_date(fixed.get("end", "")) if fixed.get("end") else ""

    periods = fixed.get("periods")
    if isinstance(periods, list) and periods:
        clean_periods = []
        for p in periods:
            if isinstance(p, dict) and p.get("start") and p.get("end"):
                clean_periods.append({"start": normalize_date(p.get("start", "")), "end": normalize_date(p.get("end", ""))})
        if clean_periods:
            fixed["periods"] = clean_periods
            fixed["start"] = clean_periods[0]["start"]
            fixed["end"] = clean_periods[-1]["end"]
            fixed["days"] = str(calc_total_days(clean_periods))
        elif "periods" in fixed:
            del fixed["periods"]

    if not fixed.get("days") and fixed.get("start") and fixed.get("end"):
        try:
            fixed["days"] = str(days_between(fixed["start"], fixed["end"]))
        except Exception:
            fixed["days"] = ""
    else:
        fixed["days"] = str(fixed.get("days", ""))

    if fixed.get("type") == TRIP_TYPE or is_material_assistance(fixed):
        fixed["return_date"] = ""
    elif fixed.get("end"):
        try:
            fixed["return_date"] = get_return_to_work_date(fixed["end"])
        except Exception:
            fixed["return_date"] = normalize_date(fixed.get("return_date", ""))
    else:
        fixed["return_date"] = normalize_date(fixed.get("return_date", ""))

    fixed["created_at"] = normalize_created_at(fixed.get("created_at"))
    if fixed.get("extended_at"):
        fixed["extended_at"] = normalize_created_at(fixed.get("extended_at"))
    return fixed


def history_unique_key(record):
    periods = record.get("periods")
    periods_text = json.dumps(periods, ensure_ascii=False, sort_keys=True) if isinstance(periods, list) else ""
    return (
        record.get("fio", ""), record.get("type", ""), record.get("start", ""),
        record.get("end", ""), record.get("days", ""), record.get("position", ""),
        record.get("project", ""), periods_text,
    )


def normalize_history_file():
    history = load_history()
    fixed_history = []
    seen = set()
    for record in history:
        if not isinstance(record, dict):
            continue
        fixed = normalize_history_record(record)
        key = history_unique_key(fixed)
        if key in seen:
            continue
        seen.add(key)
        fixed_history.append(fixed)
    save_history_full(fixed_history)
    return fixed_history


def is_same_record(a, b):
    return (
        a.get("fio") == b.get("fio") and
        a.get("type") == b.get("type") and
        a.get("start") == b.get("start") and
        a.get("end") == b.get("end") and
        a.get("created_at") == b.get("created_at")
    )


def update_history_record(old_record, new_record):
    history = load_history()
    new_history = []
    replaced = False
    for r in history:
        if not replaced and is_same_record(r, old_record):
            new_history.append(new_record)
            replaced = True
        else:
            new_history.append(r)
    if not replaced:
        new_history.append(new_record)
    save_history_full(new_history)


def save_history(d):
    return_date = ""
    if d.get("end") and not is_material_assistance(d):
        return_date = get_return_to_work_date(d["end"])
    record = {
        "fio": d.get("fio"),
        "position": d.get("pos", ""),
        "project": d.get("project", ""),
        "type": d.get("type"),
        "start": normalize_date(d.get("start", "")) if d.get("start") else "",
        "end": normalize_date(d.get("end", "")) if d.get("end") else "",
        "days": d.get("days", ""),
        "return_date": return_date,
        "created_at": now_dt().strftime("%d.%m.%Y %H:%M:%S"),
    }
    if d.get("periods"):
        record["periods"] = d.get("periods")
    append_history_record(record)


def replace_text(doc, rep):
    for p in doc.paragraphs:
        text = p.text
        for k, v in rep.items():
            text = text.replace(k, str(v))
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    for k, v in rep.items():
                        text = text.replace(k, str(v))
                    if p.runs:
                        p.runs[0].text = text
                        for r in p.runs[1:]:
                            r.text = ""


def create_doc(d):
    doc = Document(TEMPLATES[d["type"]])
    project_value = d.get("project", "")
    # В новом заявлении на рождение ребенка + 3 дня поле "Отдел" в Word оставляем пустым.
    # Проект нужен боту для истории/сообщения в группу, но в самом документе его не пишем.
    if d.get("type") == CHILD_BIRTH_3_DAYS_TYPE:
        project_value = ""
    rep = {
        "{{FIO}}": format_fio_short(d["fio"]),
        "{{POSITION}}": d.get("pos", ""),
        "{{PROJECT}}": project_value,
        "{{TODAY}}": format_date_text_without_year_word(now_dt().strftime("%d.%m.%Y")),
    }
    if "start" in d and "end" in d:
        rep["{{DATE_START}}"] = format_date_range_start(d["start"])
        rep["{{DATE_END}}"] = format_date_range_end(d["end"])
    elif "start" in d:
        rep["{{DATE_START}}"] = format_date_text(d["start"])
    replace_text(doc, rep)
    os.makedirs(READY_FOLDER, exist_ok=True)
    path = os.path.join(READY_FOLDER, f"{d['fio']}.docx")
    doc.save(path)
    return path


def month_name_capitalized(date_text):
    return month_name(date_text).capitalize()


def trip_position_text(position):
    pos = str(position or "").strip()
    if pos == "Инженер программист":
        pos = "Инженер-программист"
    if TRIP_REGION_SUFFIX.lower() not in pos.lower():
        pos = f"{pos} {TRIP_REGION_SUFFIX}".strip()
    return pos


def safe_filename(text):
    text = re.sub(r"[^А-Яа-яA-Za-z0-9_ -]", "", str(text or ""))
    text = text.strip().replace(" ", "_")
    return text or "file"


def create_trip_docs(d):
    os.makedirs(READY_FOLDER, exist_ok=True)
    start = d["start"]
    end = d["end"]
    now_text = now_dt().strftime("%d.%m.%Y")
    start_dt = datetime.strptime(start, "%d.%m.%Y")
    end_dt = datetime.strptime(end, "%d.%m.%Y")
    cur_dt = datetime.strptime(now_text, "%d.%m.%Y")
    employees_lines = [f"{emp.get('fio')} - {trip_position_text(emp.get('pos', ''))}" for emp in d.get("employees", [])]
    common_rep = {
        "{DAY_NOW}": cur_dt.strftime("%d"),
        "{MONTH_NOW}": month_name_capitalized(now_text),
        "{YEAR_NOW}": str(cur_dt.year),
        "{DAY_FROM}": start_dt.strftime("%d"),
        "{DAY_TO}": end_dt.strftime("%d"),
        "{MONTH}": month_name_capitalized(start),
        "{YEAR}": str(start_dt.year),
        "{EMPLOYEES_LIST}": "\n".join(employees_lines),
    }
    report_doc = Document(TRIP_REPORT_TEMPLATE)
    replace_text(report_doc, common_rep)
    report_path = os.path.join(READY_FOLDER, f"Командировка_рапорт_{start_dt.strftime('%d_%m_%Y')}_{end_dt.strftime('%d_%m_%Y')}.docx")
    report_doc.save(report_path)
    paths = [report_path]
    for emp in d.get("employees", []):
        cert_doc = Document(TRIP_CERTIFICATE_TEMPLATE)
        rep = {
            "{FIO}": emp.get("fio", ""),
            "{POSITION}": trip_position_text(emp.get("pos", "")),
            "{DAY_FROM}": start_dt.strftime("%d"),
            "{DAY_TO}": end_dt.strftime("%d"),
            "{MONTH}": month_name_capitalized(start),
            "{YEAR}": str(start_dt.year),
        }
        replace_text(cert_doc, rep)
        cert_path = os.path.join(READY_FOLDER, f"Командировочное_удостоверение_{safe_filename(emp.get('fio', ''))}.docx")
        cert_doc.save(cert_path)
        paths.append(cert_path)
    return paths


def save_trip_history(d):
    history = load_history()
    group_id = d.get("trip_group_id") or now_dt().strftime("%Y%m%d%H%M%S") + "_" + str(uuid.uuid4())[:8]
    created_at = now_dt().strftime("%d.%m.%Y %H:%M:%S")
    try:
        days = str(days_between(d.get("start", ""), d.get("end", "")))
    except Exception:
        days = ""
    trip_employees = [{"fio": emp.get("fio", ""), "position": trip_position_text(emp.get("pos", ""))} for emp in d.get("employees", [])]
    for emp in d.get("employees", []):
        history.append({
            "fio": emp.get("fio", ""),
            "position": trip_position_text(emp.get("pos", "")),
            "project": "",
            "type": TRIP_TYPE,
            "start": normalize_date(d.get("start", "")),
            "end": normalize_date(d.get("end", "")),
            "days": days,
            "return_date": "",
            "created_at": created_at,
            "trip_group_id": group_id,
            "trip_employees": trip_employees,
        })
    save_history_full(history)


def get_trip_groups(limit=10):
    history = load_history()
    groups = {}
    for r in history:
        if r.get("type") != TRIP_TYPE:
            continue
        group_id = r.get("trip_group_id") or f"{r.get('created_at','')}_{r.get('start','')}_{r.get('end','')}"
        if group_id not in groups:
            groups[group_id] = {"trip_group_id": group_id, "start": r.get("start", ""), "end": r.get("end", ""), "created_at": r.get("created_at", ""), "employees": []}
        if r.get("trip_employees"):
            groups[group_id]["employees"] = r.get("trip_employees", [])
        else:
            groups[group_id]["employees"].append({"fio": r.get("fio", ""), "position": r.get("position", "")})
    result = list(groups.values())
    result.sort(key=lambda x: normalize_created_at(x.get("created_at", "")), reverse=True)
    clean = []
    for g in result:
        seen = set()
        employees = []
        for emp in g.get("employees", []):
            fio = emp.get("fio", "")
            if not fio or fio in seen:
                continue
            seen.add(fio)
            employees.append(emp)
        g["employees"] = employees
        clean.append(g)
    return clean[:limit]


def trip_repeat_keyboard(groups):
    return make_keyboard([str(i) for i in range(1, len(groups) + 1)] + ["🏠 Старт"], cols=3)


def finish_and_send(chat_id):
    path = create_doc(data[chat_id])
    save_history(data[chat_id])
    return path


async def send_group_message(text):
    try:
        await bot.send_message(chat_id=GROUP_CHAT_ID, text=text)
    except Exception as e:
        print("GROUP SEND ERROR:", e)


async def notify_application_created(d):
    return_date = ""
    if d.get("end"):
        return_date = get_return_to_work_date(d["end"])
    msg = (
        f"📄 ARIZA TAYYORLANDI\n\n"
        f"👤 Xodim: {group_value(d.get('fio'))}\n"
        f"💼 Lavozim: {position_group_value(d.get('pos', ''))}\n"
        f"📌 Loyiha: {group_value(d.get('project', ''))}\n\n"
        f"📝 Ariza turi: {type_uz(d.get('type'))}\n"
    )
    if d.get("start") and d.get("end"):
        msg += f"📅 Muddat: {d.get('start')} — {d.get('end')}\n"
    elif d.get("start"):
        msg += f"📅 Sana: {d.get('start')}\n"
    if d.get("days"):
        msg += f"⏳ Davomiyligi: {d.get('days')} kun\n"
    if d.get("type") == "📌 Оставшийся отпуск":
        msg += "\nℹ️ Xodim mehnat ta’tilining qolgan qismi bo‘yicha ariza yozdi.\n"
    if return_date:
        msg += f"📌 Ishga chiqish sanasi: {return_date}\n"
    msg += "\n✅ Ariza bot orqali shakllantirildi."
    await send_group_message(msg)


async def notify_sick_created(d):
    """
    Уведомление в группу при создании новой записи больничного.
    История уже сохраняется через save_history(), поэтому запись не пропадает при Railway deploy.
    """
    return_date = ""
    if d.get("end"):
        return_date = get_return_to_work_date(d["end"])

    pos_text = d.get("pos") or display_position(d.get("fio", ""), d.get("position", ""))

    msg = (
        f"🏥 KASALLIK TA’TILI QAYD ETILDI\n\n"
        f"👤 Xodim: {group_value(d.get('fio'))}\n"
        f"💼 Lavozim: {position_group_value(pos_text)}\n"
        f"📌 Loyiha: {group_value(d.get('project', ''))}\n\n"
        f"📝 Ta’til turi: {type_uz(d.get('type'))}\n"
    )

    if d.get("start") and d.get("end"):
        msg += f"📅 Muddat: {d.get('start')} — {d.get('end')}\n"
    elif d.get("start"):
        msg += f"📅 Sana: {d.get('start')}\n"

    if d.get("days"):
        msg += f"⏳ Davomiyligi: {d.get('days')} kun\n"

    if return_date:
        msg += f"📌 Ishga chiqish sanasi: {return_date}\n"

    msg += "\n✅ Ma’lumot bot orqali kiritildi."
    await send_group_message(msg)


async def notify_sick_extended(old_record, new_record, added_start, added_end):
    msg = (
        f"🏥 KASALLIK TA’TILI UZAYTIRILDI\n\n"
        f"👤 Xodim: {group_value(new_record.get('fio'))}\n"
        f"💼 Lavozim: {position_group_value(new_record.get('position', ''))}\n"
        f"📌 Loyiha: {group_value(new_record.get('project', ''))}\n\n"
        f"📅 Avvalgi muddat: {old_record.get('start')} — {old_record.get('end')}\n"
        f"➕ Qo‘shilgan muddat: {added_start} — {added_end}\n\n"
        f"📋 Umumiy davrlar:\n{format_periods(get_periods_from_record(new_record))}\n\n"
        f"⏳ Jami davomiyligi: {new_record.get('days')} kun\n"
        f"📌 Ishga chiqish sanasi: {new_record.get('return_date', '')}\n\n"
        f"ℹ️ Agar oxirgi kun dam olish yoki bayram kuniga to‘g‘ri kelsa, ishga chiqish sanasi keyingi ish kuniga o‘tkazildi."
    )
    await send_group_message(msg)


async def notify_bs_extended(old_record, new_record, added_start, added_end):
    msg = (
        f"📝 ISH HAQI SAQLANMAGAN TA’TIL UZAYTIRILDI\n\n"
        f"👤 Xodim: {group_value(new_record.get('fio'))}\n"
        f"💼 Lavozim: {position_group_value(new_record.get('position', ''))}\n"
        f"📌 Loyiha: {group_value(new_record.get('project', ''))}\n\n"
        f"📅 Avvalgi muddat: {old_record.get('start')} — {old_record.get('end')}\n"
        f"➕ Qo‘shilgan muddat: {added_start} — {added_end}\n\n"
        f"📋 Umumiy davrlar:\n{format_periods(get_periods_from_record(new_record))}\n\n"
        f"⏳ Jami davomiyligi: {new_record.get('days')} kun\n"
        f"📌 Ishga chiqish sanasi: {new_record.get('return_date', '')}\n\n"
        f"ℹ️ Agar oxirgi kun dam olish yoki bayram kuniga to‘g‘ri kelsa, ishga chiqish sanasi keyingi ish kuniga o‘tkazildi."
    )
    await send_group_message(msg)


def employee_keyboard(employees):
    return make_keyboard(employees + ["🏠 Старт"], cols=2)


def is_active_record(record, today):
    try:
        start = parse_date_or_none(record.get("start", ""))
        end = parse_date_or_none(record.get("end", ""))

        if not start or not end:
            return False

        # Обычная логика отчета: показываем с начала до конца отсутствия.
        if start <= today <= end:
            return True

        # Особая логика дня выхода:
        # если сегодня return_date, сотрудник виден в отчете только ДО отправки
        # группового сообщения о выходе на работу. После отправки исчезает.
        today_text = today.strftime("%d.%m.%Y")
        return_date = normalize_date(record.get("return_date", ""))

        if return_date == today_text:
            sent = set(load_sent_reminders())
            for remind_time in REMIND_TIMES:
                key = f"return_{today_text}_{remind_time}_{record.get('fio')}_{record.get('type')}"
                if key in sent:
                    return False
            return True

        return False

    except Exception:
        return False


def build_report():
    today = now_dt().date()
    history = load_history()
    groups = {
        "🌴 В отпуске": [],
        "📝 В БС": [],
        "📚 В учебном отпуске": [],
        "🏥 На больничном": [],
        "✈️ В командировке": [],
    }

    for r in history:
        if not isinstance(r, dict):
            continue
        if is_material_assistance(r):
            continue
        if not is_active_record(r, today):
            continue

        t = r.get("type", "")
        if t in ["🌴 Полный отпуск", "🧩 Часть отпуска", "📌 Оставшийся отпуск", CHILD_BIRTH_3_DAYS_TYPE]:
            groups["🌴 В отпуске"].append(r)
        elif t in ["📝 БС с периода по период", "📅 БС на один день"]:
            groups["📝 В БС"].append(r)
        elif t == "📚 Учебный отпуск":
            groups["📚 В учебном отпуске"].append(r)
        elif t == SICK_LEAVE_TYPE:
            groups["🏥 На больничном"].append(r)
        elif t == TRIP_TYPE:
            groups["✈️ В командировке"].append(r)

    msg = f"📊 Отчет на {now_dt().strftime('%d.%m.%Y')}\n\n"
    for title, records in groups.items():
        msg += f"{title}:\n"
        if not records:
            msg += "Нет сотрудников\n\n"
            continue
        for i, r in enumerate(records, 1):
            msg += f"{i}. {r.get('fio')}\n"
            pos_text = display_position(r.get("fio", ""), r.get("position", ""))
            if pos_text:
                msg += f"   Должность: {pos_text}\n"
            if r.get("type") != TRIP_TYPE:
                msg += f"   Проект: {r.get('project', '')}\n"
            msg += f"   Тип: {r.get('type')}\n"
            if r.get("periods"):
                msg += "   Периоды:\n"
                for p in get_periods_from_record(r):
                    msg += f"   - {p.get('start')} по {p.get('end')}\n"
            else:
                msg += f"   С: {r.get('start')} по {r.get('end')}\n"
            if r.get("days"):
                msg += f"   Дней: {r.get('days')}\n"
            if r.get("return_date"):
                msg += f"   Выход: {normalize_date(r.get('return_date'))}\n"
            msg += "\n"
    return msg


def time_is_due(now_time, target_time):
    try:
        now_value = datetime.strptime(now_time, "%H:%M").time()
        target_value = datetime.strptime(target_time, "%H:%M").time()
        return now_value >= target_value
    except Exception:
        return False


async def reminder_loop():
    """
    Экономный цикл напоминаний.

    Раньше бот каждые 30 секунд читал history_records и sent_reminders из Neon.
    Это быстро расходовало compute time в Neon.

    Теперь бот НЕ дергает базу постоянно:
    - в обычное время просто спит;
    - базу читает только когда наступило время напоминаний;
    - каждое окно напоминаний обрабатывает один раз за день;
    - если бот перезапустился после нужного времени, он все равно проверит напоминания,
      но снова не отправит дубли, потому что reserve_sent_reminder() проверяет sent_reminders.
    """
    print("REMINDER LOOP STARTED")
    processed_windows = set()

    while True:
        sleep_seconds = 60
        try:
            now = now_dt()
            now_date = now.strftime("%d.%m.%Y")
            now_time = now.strftime("%H:%M")
            tomorrow_date = (now + timedelta(days=1)).strftime("%d.%m.%Y")

            due_events = []

            # Напоминание о начале отпуска/БС/учебного отпуска завтра.
            # Проверяем только один раз в день после START_REMIND_TIME.
            start_window_key = f"start_window_{now_date}_{START_REMIND_TIME}"
            if time_is_due(now_time, START_REMIND_TIME) and start_window_key not in processed_windows:
                due_events.append(("start", START_REMIND_TIME, start_window_key))

            # Напоминание о выходе на работу сегодня.
            # Проверяем один раз в день для каждого времени из REMIND_TIMES.
            for remind_time in REMIND_TIMES:
                return_window_key = f"return_window_{now_date}_{remind_time}"
                if time_is_due(now_time, remind_time) and return_window_key not in processed_windows:
                    due_events.append(("return", remind_time, return_window_key))

            # Если сейчас нет нужного окна — не трогаем Neon.
            if not due_events:
                await asyncio.sleep(sleep_seconds)
                continue

            # ВАЖНО: базу читаем только здесь, когда реально есть что проверить.
            history = load_history()
            sent = set(load_sent_reminders())

            for event_type, remind_time, window_key in due_events:
                if event_type == "start":
                    for r in history:
                        if not isinstance(r, dict):
                            continue
                        if r.get("type") == TRIP_TYPE or is_material_assistance(r):
                            continue
                        if normalize_date(r.get("start", "")) == tomorrow_date:
                            key = f"start_{now_date}_{r.get('fio')}_{r.get('start')}_{r.get('type')}"
                            if key in sent:
                                continue
                            if not reserve_sent_reminder(key):
                                sent.add(key)
                                continue
                            sent.add(key)
                            msg = (
                                f"🔔 ERTAGA TA’TIL BOSHLANADI\n\n"
                                f"👤 Xodim: {group_value(r.get('fio'))}\n"
                                f"📌 Loyiha: {group_value(r.get('project', ''))}\n\n"
                                f"📝 Ta’til turi: {type_uz(r.get('type'))}\n"
                                f"📅 Muddat: {r.get('start')} — {r.get('end')}\n"
                                f"⏳ Davomiyligi: {r.get('days')} kun\n"
                                f"📌 Ishga chiqish sanasi: {normalize_date(r.get('return_date', ''))}\n\n"
                                f"ℹ️ Xodim ertadan boshlab {start_phrase_uz(r.get('type'))}."
                            )
                            await send_group_message(msg)
                            print("START REMINDER SENT:", key)

                elif event_type == "return":
                    for r in history:
                        if not isinstance(r, dict):
                            continue
                        if r.get("type") == TRIP_TYPE or is_material_assistance(r):
                            continue

                        if normalize_date(r.get("return_date", "")) == now_date:
                            key = f"return_{now_date}_{remind_time}_{r.get('fio')}_{r.get('type')}"
                            if key in sent:
                                continue
                            if not reserve_sent_reminder(key):
                                sent.add(key)
                                continue
                            sent.add(key)
                            msg = (
                                f"✅ BUGUN ISHGA CHIQADI\n\n"
                                f"👤 Xodim: {group_value(r.get('fio'))}\n"
                                f"📌 Loyiha: {group_value(r.get('project', ''))}\n\n"
                                f"📝 Ta’til turi: {type_uz(r.get('type'))}\n"
                                f"📅 Ishga chiqish sanasi: {normalize_date(r.get('return_date', ''))}\n\n"
                                f"ℹ️ Xodim bugundan ish faoliyatini davom ettiradi."
                            )
                            await send_group_message(msg)
                            print("RETURN REMINDER SENT:", key)

                processed_windows.add(window_key)

        except Exception as e:
            print("REMINDER LOOP ERROR:", e)

        await asyncio.sleep(sleep_seconds)


menu = make_keyboard([
    "📄 Создать заявление",
    "✈️ Командировка",
    "➕ Добавить запись вручную",
    "🏥 Больничный",
    "📊 Отчет",
    "🗑 Удалить запись",
    "📜 История",
    "🏠 Старт",
], cols=2)

report_only_menu = make_keyboard(["📊 Отчет", "🏠 Старт"], cols=2)
history_menu = make_keyboard(["🔍 Поиск сотрудника", "📋 Полный список сотрудников", "🏠 Старт"], cols=2)
history_action_menu = make_keyboard(["✏️ Изменить дату выхода", "🏠 Старт"], cols=1)
confirm_return_change_menu = make_keyboard(["✅ Да", "❌ Нет"], cols=2)
pos_menu = make_keyboard(["Инженер программист", "Программист", "🏠 Старт"], cols=2)
confirm_delete_menu = make_keyboard(["✅ Да, удалить", "❌ Нет, отменить"], cols=2)
manual_type_menu = make_keyboard(MANUAL_TYPES + ["🏠 Старт"], cols=2)
saved_profile_menu = make_keyboard(["✅ Да, использовать", "✏️ Изменить", "🏠 Старт"], cols=1)
report_export_menu = make_keyboard(["📥 Excel за этот месяц", "📥 Excel за прошлый месяц", "💰 Зарплаты", "🏠 Старт"], cols=1)
salary_menu = make_keyboard(["🔍 Найти сотрудника", "📋 Весь список", "🏠 Старт"], cols=1)
overlap_menu = make_keyboard(["✅ Всё равно сохранить", "✏️ Изменить дату", "❌ Отмена"], cols=1)


def find_active_sick_records_by_fio(fio):
    today = now_dt().date()
    return [r for r in load_history() if r.get("fio") == fio and r.get("type") == SICK_LEAVE_TYPE and is_active_record(r, today)]


def find_active_sick_records_by_text(text):
    today = now_dt().date()
    return [r for r in load_history() if r.get("type") == SICK_LEAVE_TYPE and fio_startswith_text(r.get("fio", ""), text) and is_active_record(r, today)]


def find_active_bs_records_by_fio(fio):
    today = now_dt().date()
    return [r for r in load_history() if r.get("fio") == fio and r.get("type") in BS_LEAVE_TYPES and is_active_record(r, today)]


def find_previous_part_leave_records_by_fio(fio):
    return [r for r in load_history() if r.get("fio") == fio and r.get("type") in ["🧩 Часть отпуска", "📌 Оставшийся отпуск"]]


def format_leave_records_for_message(records):
    msg = ""
    for i, r in enumerate(records, 1):
        msg += f"{i}. {r.get('type')} — {r.get('start')} по {r.get('end')}"
        if r.get("days"):
            msg += f" ({r.get('days')} kun)"
        msg += "\n"
    return msg


def can_change_return_date(record):
    if not isinstance(record, dict):
        return False
    if record.get("type") == TRIP_TYPE or is_material_assistance(record):
        return False
    return bool(record.get("return_date"))


def choose_return_change_record(records):
    records = [r for r in records if isinstance(r, dict)]
    if not records:
        return None

    today = now_dt().date()

    def sort_key(record):
        start = parse_date_or_none(record.get("start", ""))
        end = parse_date_or_none(record.get("end", ""))
        return_date = parse_date_or_none(record.get("return_date", ""))
        created = normalize_created_at(record.get("created_at", ""))

        # Сначала берем текущую/актуальную запись: отпуск идет сейчас
        # или дата выхода еще впереди. Старые отпуска не мешают.
        active_score = 0
        if start and return_date and start <= today <= return_date:
            active_score = 2
        elif return_date and return_date >= today:
            active_score = 1

        return (active_score, return_date or end or start or today, created)

    records.sort(key=sort_key, reverse=True)
    return records[0]


def return_change_prompt(record):
    return (
        f"Введите новую дату выхода на работу ДД.ММ.ГГГГ\n\n"
        f"👤 {record.get('fio')}\n"
        f"Тип: {record.get('type')}\n"
        f"С: {record.get('start')} по {record.get('end')}\n"
        f"Текущий выход: {normalize_date(record.get('return_date', ''))}"
    )


async def notify_return_date_changed(old_record, new_record, reason="bayram / qo‘shimcha dam olish kuni"):
    msg = (
        f"📢 ISHGA CHIQISH SANASI O‘ZGARTIRILDI\n\n"
        f"👤 Xodim: {group_value(new_record.get('fio'))}\n"
        f"💼 Lavozim: {position_group_value(new_record.get('position', ''))}\n"
        f"📌 Loyiha: {group_value(new_record.get('project', ''))}\n\n"
        f"📝 Ta’til turi: {type_uz(new_record.get('type'))}\n"
        f"📅 Ta’til muddati: {new_record.get('start')} — {new_record.get('end')}\n"
        f"📌 Avvalgi ishga chiqish sanasi: {normalize_date(old_record.get('return_date', ''))}\n"
        f"✅ Yangi ishga chiqish sanasi: {normalize_date(new_record.get('return_date', ''))}\n\n"
        f"ℹ️ Sabab: {reason}."
    )
    await send_group_message(msg)


def build_extended_record(old_record, new_start, new_end, force_type=None):
    old_periods = get_periods_from_record(old_record)
    periods = old_periods + [{"start": new_start, "end": new_end}]
    new_record = dict(old_record)
    if force_type:
        new_record["type"] = force_type
    new_record["start"] = periods[0].get("start", "")
    new_record["end"] = new_end
    new_record["days"] = str(calc_total_days(periods))
    new_record["return_date"] = get_return_to_work_date(new_end)
    new_record["periods"] = periods
    new_record["extended_at"] = now_dt().strftime("%d.%m.%Y %H:%M:%S")
    return new_record


def clean_position_for_profile(position):
    pos = str(position or "").strip()
    pos = pos.replace(TRIP_REGION_SUFFIX, "").replace("Инженер-программист", "Инженер программист").strip()
    if pos in ["Инженер программист", "Программист"]:
        return pos
    return pos


def get_employee_last_profile(fio):
    records = [r for r in load_history() if isinstance(r, dict) and r.get("fio") == fio]
    records.sort(key=lambda r: normalize_created_at(r.get("created_at", "")), reverse=True)

    # Должность берем из нового файла employees/salary.docx, чтобы она всегда была актуальной.
    last_pos = clean_position_for_profile(get_position_from_salary(fio))
    last_project = ""

    # Проект берем из последней истории сотрудника.
    for r in records:
        if not last_project:
            project = str(r.get("project", "")).strip()
            if project:
                last_project = project
        if not last_pos:
            pos = clean_position_for_profile(r.get("position", ""))
            if pos:
                last_pos = pos
        if last_pos and last_project:
            break
    if not last_pos and not last_project:
        return None
    return {"pos": last_pos, "project": last_project}


def profile_message(profile):
    msg = "Найдены сохраненные данные:\n\n"
    msg += f"Должность: {profile.get('pos') or 'не указана'}\n"
    msg += f"Проект: {profile.get('project') or 'не указан'}\n\n"
    msg += "Использовать эти данные?"
    return msg


def records_overlap(start1, end1, start2, end2):
    return start1 <= end2 and start2 <= end1


def find_date_overlaps(fio, start_text, end_text):
    start = parse_date_or_none(start_text)
    end = parse_date_or_none(end_text)
    if not start or not end:
        return []
    result = []
    for r in load_history():
        if not isinstance(r, dict):
            continue
        if r.get("fio") != fio:
            continue
        if r.get("type") not in [
            "🌴 Полный отпуск", "🧩 Часть отпуска", "📌 Оставшийся отпуск",
            "📚 Учебный отпуск", "📝 БС с периода по период", "📅 БС на один день",
            SICK_LEAVE_TYPE, CHILD_BIRTH_3_DAYS_TYPE, TRIP_TYPE,
        ]:
            continue
        periods = get_periods_from_record(r)
        if not periods and r.get("start") and r.get("end"):
            periods = [{"start": r.get("start"), "end": r.get("end")}]
        for p in periods:
            old_start = parse_date_or_none(p.get("start", ""))
            old_end = parse_date_or_none(p.get("end", ""))
            if old_start and old_end and records_overlap(start, end, old_start, old_end):
                result.append(r)
                break
    return result


def format_overlap_warning(overlaps, new_start, new_end):
    msg = "⚠️ У этого сотрудника уже есть запись на эти даты:\n\n"
    for i, r in enumerate(overlaps[:5], 1):
        msg += f"{i}. {r.get('type')}\n"
        msg += f"   С: {r.get('start')} по {r.get('end')}\n"
        if r.get("project"):
            msg += f"   Проект: {r.get('project')}\n"
        msg += "\n"
    msg += f"Новая запись: {new_start} — {new_end}\n\n"
    msg += "Всё равно сохранить?"
    return msg


def need_overlap_check(chat_id):
    return not data.get(chat_id, {}).get("ignore_overlap")


async def check_overlap_or_continue(m, chat_id, action_name):
    d = data.get(chat_id, {})
    if need_overlap_check(chat_id) and d.get("fio") and d.get("start") and d.get("end"):
        overlaps = find_date_overlaps(d.get("fio"), d.get("start"), d.get("end"))
        if overlaps:
            d["pending_action"] = action_name
            state[chat_id] = "overlap_confirm"
            await m.answer(format_overlap_warning(overlaps, d.get("start"), d.get("end")), reply_markup=overlap_menu)
            return False
    return True


async def finalize_doc_action(m, chat_id):
    path = finish_and_send(chat_id)
    await m.answer_document(FSInputFile(path))

    # Матпомощь (свадьба/ребенок/смерть родственника) не отправляем в группу:
    # это только формирование документа и история, не кадровое уведомление.
    if not is_material_assistance(data[chat_id]):
        await notify_application_created(data[chat_id])

    data[chat_id].pop("ignore_overlap", None)
    data[chat_id].pop("pending_action", None)
    state[chat_id] = "menu"
    await m.answer("Готово", reply_markup=menu)


async def finalize_manual_action(m, chat_id):
    save_history(data[chat_id])
    if data[chat_id].get("type") == SICK_LEAVE_TYPE:
        await notify_sick_created(data[chat_id])
    data[chat_id].pop("ignore_overlap", None)
    data[chat_id].pop("pending_action", None)
    state[chat_id] = "menu"
    await m.answer("Запись вручную добавлена ✅", reply_markup=menu)


async def finalize_sick_action(m, chat_id):
    save_history(data[chat_id])
    await notify_sick_created(data[chat_id])
    data[chat_id].pop("ignore_overlap", None)
    data[chat_id].pop("pending_action", None)
    state[chat_id] = "menu"
    await m.answer("Больничный сохранен в историю и отправлен в группу ✅", reply_markup=menu)


def find_trip_overlaps(d):
    overlaps = []
    for emp in d.get("employees", []):
        emp_overlaps = find_date_overlaps(emp.get("fio"), d.get("start"), d.get("end"))
        for r in emp_overlaps:
            overlaps.append((emp.get("fio"), r))
    return overlaps


def format_trip_overlap_warning(overlaps, new_start, new_end):
    msg = "⚠️ По командировке найдены пересечения дат:\n\n"
    for i, item in enumerate(overlaps[:8], 1):
        fio, r = item
        msg += f"{i}. {fio} — {r.get('type')}\n"
        msg += f"   С: {r.get('start')} по {r.get('end')}\n\n"
    msg += f"Новая командировка: {new_start} — {new_end}\n\n"
    msg += "Всё равно сохранить?"
    return msg


async def finalize_trip_action(m, chat_id):
    paths = create_trip_docs(data[chat_id])
    save_trip_history(data[chat_id])
    for path in paths:
        await m.answer_document(FSInputFile(path))
    data[chat_id].pop("ignore_overlap", None)
    data[chat_id].pop("pending_action", None)
    state[chat_id] = "menu"
    await m.answer("Командировка готова ✅ и сохранена в историю", reply_markup=menu)


async def run_pending_action(m, chat_id):
    action = data.get(chat_id, {}).get("pending_action")
    if action == "doc":
        await finalize_doc_action(m, chat_id)
    elif action == "manual":
        await finalize_manual_action(m, chat_id)
    elif action == "sick":
        await finalize_sick_action(m, chat_id)
    elif action == "trip":
        await finalize_trip_action(m, chat_id)
    else:
        state[chat_id] = "menu"
        await m.answer("Действие отменено.", reply_markup=menu)


def month_bounds(offset=0):
    today = now_dt().date()
    year = today.year
    month = today.month + offset
    while month < 1:
        month += 12
        year -= 1
    while month > 12:
        month -= 12
        year += 1
    start = datetime(year, month, 1).date()
    end = datetime(year, month, monthrange(year, month)[1]).date()
    return start, end


def create_excel_report(offset=0):
    if Workbook is None:
        return None
    start_month, end_month = month_bounds(offset)
    records = []
    for r in load_history():
        if not isinstance(r, dict):
            continue
        start = parse_date_or_none(r.get("start", ""))
        end = parse_date_or_none(r.get("end", ""))
        if not start or not end:
            continue
        if records_overlap(start, end, start_month, end_month):
            records.append(r)
    records.sort(key=lambda r: (parse_date_or_none(r.get("start", "")) or start_month, r.get("fio", "")))

    wb = Workbook()
    ws = wb.active
    ws.title = "Отчет"
    title = f"Отчет за {start_month.strftime('%m.%Y')}"
    ws.append([title])
    ws.append([])
    headers = ["№", "ФИО", "Должность", "Проект", "Тип", "Начало", "Конец", "Дней", "Выход", "Создан"]
    ws.append(headers)
    for cell in ws[3]:
        if Font:
            cell.font = Font(bold=True)
        if Alignment:
            cell.alignment = Alignment(horizontal="center")
    for i, r in enumerate(records, 1):
        ws.append([
            i, r.get("fio", ""), display_position(r.get("fio", ""), r.get("position", "")), r.get("project", ""), r.get("type", ""),
            r.get("start", ""), r.get("end", ""), r.get("days", ""), r.get("return_date", ""), r.get("created_at", ""),
        ])
    widths = [6, 32, 28, 24, 26, 14, 14, 10, 14, 22]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + idx)].width = width
    os.makedirs(READY_FOLDER, exist_ok=True)
    path = os.path.join(READY_FOLDER, f"office_report_{start_month.strftime('%Y_%m')}.xlsx")
    wb.save(path)
    return path


def normalize_salary_number(value):
    text = str(value or "").strip().replace(",", ".")
    if not text:
        return ""
    try:
        number = float(text)
        amount = int(number * 1000000)
        return f"{amount:,}".replace(",", " ") + " сум"
    except Exception:
        return text



def load_salary_records():
    global SALARY_RECORDS_CACHE

    if not os.path.exists(SALARY_FILE):
        SALARY_RECORDS_CACHE = {"mtime": None, "records": []}
        return []

    try:
        current_mtime = os.path.getmtime(SALARY_FILE)
    except Exception:
        current_mtime = None

    if (
        SALARY_RECORDS_CACHE.get("records") is not None
        and SALARY_RECORDS_CACHE.get("mtime") == current_mtime
    ):
        return [dict(r) for r in SALARY_RECORDS_CACHE.get("records", [])]

    try:
        doc = Document(SALARY_FILE)
        records = []

        # Новый файл employees/salary.docx обычно сделан таблицей:
        # № | ФИО | Должность | сумма. Поэтому сначала читаем таблицы.
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) < 4:
                    continue
                number, fio, position, salary_short = cells[0], cells[1], cells[2], cells[3]
                if not str(number).strip().isdigit() or not fio:
                    continue
                records.append({
                    "fio": fio.strip(),
                    "position": position.strip(),
                    "salary_short": salary_short.strip(),
                    "salary": normalize_salary_number(salary_short),
                })

        if records:
            SALARY_RECORDS_CACHE = {"mtime": current_mtime, "records": records}
            return [dict(r) for r in records]

        # Резервный вариант: если файл будет не таблицей, а обычным текстом
        # в формате: номер / ФИО / должность / сумма.
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        i = 0
        while i < len(lines):
            if lines[i].isdigit() and i + 3 < len(lines):
                fio = lines[i + 1].strip()
                position = lines[i + 2].strip()
                salary_short = lines[i + 3].strip()
                records.append({
                    "fio": fio,
                    "position": position,
                    "salary_short": salary_short,
                    "salary": normalize_salary_number(salary_short),
                })
                i += 4
            else:
                i += 1

        SALARY_RECORDS_CACHE = {"mtime": current_mtime, "records": records}
        return [dict(r) for r in records]
    except Exception as e:
        print("SALARY FILE READ ERROR:", e)
        SALARY_RECORDS_CACHE = {"mtime": current_mtime, "records": []}
        return []

def search_salary_records_by_text(text):
    q = str(text or "").lower().strip()
    if len(q) < 2:
        return None
    return [r for r in load_salary_records() if any(part.lower().startswith(q) for part in r.get("fio", "").split())]


def get_salary_record_by_fio(fio):
    fio = str(fio or "").strip()
    if not fio:
        return None
    for r in load_salary_records():
        if r.get("fio") == fio:
            return r
    return None


def get_position_from_salary(fio):
    record = get_salary_record_by_fio(fio)
    return record.get("position", "") if record else ""


def display_position(fio, saved_position=""):
    # В отчетах показываем должность из актуального файла salary.docx.
    return get_position_from_salary(fio) or str(saved_position or "").strip()


def salary_record_text(record, number=None):
    prefix = f"{number}. " if number is not None else ""
    return (
        f"{prefix}{record.get('fio', '')}\n"
        f"   Должность: {record.get('position', '')}\n"
        f"   Зарплата: {record.get('salary', '')}\n"
    )


async def send_long_message(message, text, reply_markup=None):
    max_len = 3800
    text = str(text or "")
    parts = []
    while len(text) > max_len:
        cut = text.rfind("\n", 0, max_len)
        if cut == -1:
            cut = max_len
        parts.append(text[:cut].strip())
        text = text[cut:].strip()
    if text:
        parts.append(text)
    if not parts:
        parts = [""]
    for i, part in enumerate(parts):
        await message.answer(part, reply_markup=reply_markup if i == len(parts) - 1 else None)


@dp.message(Command("start"))
async def start(m: Message):
    chat_id = m.chat.id
    if chat_id == GROUP_CHAT_ID:
        return
    if not is_allowed(chat_id):
        await m.answer("⛔ Доступ запрещен")
        return
    state[chat_id] = "menu"
    await m.answer("Меню", reply_markup=get_menu(chat_id))


@dp.message()
async def handler(m: Message):
    chat_id = m.chat.id
    text = m.text or ""

    if chat_id == GROUP_CHAT_ID:
        return
    if not is_allowed(chat_id):
        await m.answer("⛔ Доступ запрещен")
        return

    if text == "🏠 Старт":
        state[chat_id] = "menu"
        await m.answer("Меню", reply_markup=get_menu(chat_id))
        return

    if text == "📊 Отчет":
        await m.answer(build_report(), reply_markup=report_export_menu)
        return

    if text in ["📥 Excel за этот месяц", "📥 Excel за прошлый месяц"]:
        offset = 0 if text == "📥 Excel за этот месяц" else -1
        path = create_excel_report(offset)
        if not path:
            await m.answer("Для Excel-отчета нужно добавить openpyxl в requirements.txt: openpyxl", reply_markup=report_export_menu)
            return
        await m.answer_document(FSInputFile(path), reply_markup=report_export_menu)
        return

    if text == "💰 Зарплаты":
        state[chat_id] = "salary_menu"
        await m.answer("Выбери действие по зарплатам:", reply_markup=salary_menu)
        return

    if state.get(chat_id) == "salary_menu":
        if text == "🔍 Найти сотрудника":
            state[chat_id] = "salary_search"
            await m.answer("Напиши минимум 2 буквы фамилии или имени", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            return
        if text == "📋 Весь список":
            records = load_salary_records()
            if not records:
                await m.answer(f"Файл с зарплатами не найден или пустой: {SALARY_FILE}", reply_markup=salary_menu)
                return
            msg = "💰 Зарплаты сотрудников:\n\n"
            for i, record in enumerate(records, 1):
                msg += salary_record_text(record, i) + "\n"
            await send_long_message(m, msg, reply_markup=salary_menu)
            return
        await m.answer("Выбери действие кнопкой.", reply_markup=salary_menu)
        return

    if state.get(chat_id) == "salary_search":
        found = search_salary_records_by_text(text)
        if found is None:
            await m.answer("Напиши минимум 2 буквы фамилии или имени.", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            return
        if not found:
            await m.answer("Сотрудник не найден. Напиши другую часть ФИО.", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            return
        temp_search[chat_id] = found
        state[chat_id] = "salary_choose"
        await m.answer("Выбери сотрудника:", reply_markup=employee_keyboard([r.get("fio", "") for r in found]))
        return

    if state.get(chat_id) == "salary_choose":
        records = temp_search.get(chat_id, [])
        selected = None
        for record in records:
            if record.get("fio") == text:
                selected = record
                break
        if not selected:
            await m.answer("Выбери сотрудника только из списка кнопок.", reply_markup=employee_keyboard([r.get("fio", "") for r in records]))
            return
        msg = "💰 Информация по зарплате:\n\n" + salary_record_text(selected)
        state[chat_id] = "salary_menu"
        await m.answer(msg, reply_markup=salary_menu)
        return

    if state.get(chat_id) == "overlap_confirm":
        if text == "✅ Всё равно сохранить":
            data[chat_id]["ignore_overlap"] = True
            await run_pending_action(m, chat_id)
            return
        if text == "✏️ Изменить дату":
            action = data.get(chat_id, {}).get("pending_action")
            data[chat_id].pop("start", None)
            data[chat_id].pop("end", None)
            data[chat_id].pop("days", None)
            data[chat_id].pop("ignore_overlap", None)
            if action == "trip":
                state[chat_id] = "trip_start_date"
                await m.answer("Введи дату начала командировки ДД.ММ.ГГГГ", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            elif action == "sick":
                state[chat_id] = "sick_start_date"
                await m.answer("Введи дату начала больничного ДД.ММ.ГГГГ", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            elif action == "manual":
                state[chat_id] = "manual_start"
                await m.answer("Введи дату начала ДД.ММ.ГГГГ", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            else:
                state[chat_id] = "date"
                await m.answer("Введи дату начала ДД.ММ.ГГГГ", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            return
        if text == "❌ Отмена":
            state[chat_id] = "menu"
            await m.answer("Сохранение отменено.", reply_markup=get_menu(chat_id))
            return
        await m.answer("Выбери действие кнопкой.", reply_markup=overlap_menu)
        return

    if is_report_only(chat_id):
        await m.answer("У вас доступ только к отчету.", reply_markup=report_only_menu)
        return

    # ================== КОМАНДИРОВКА ==================
    if text == "✈️ Командировка":
        data[chat_id] = {"employees": []}
        state[chat_id] = "trip_main_menu"
        await m.answer("Выбери действие по командировке:", reply_markup=make_keyboard(["🆕 Новая командировка", "🔁 Повторить командировку", "🏠 Старт"], cols=1))
        return

    if state.get(chat_id) == "trip_main_menu":
        if text == "🆕 Новая командировка":
            data[chat_id] = {"employees": []}
            state[chat_id] = "trip_count"
            await m.answer("Сколько сотрудников едет в командировку? Напиши цифрой, например: 5", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            return
        if text == "🔁 Повторить командировку":
            groups = get_trip_groups(limit=10)
            if not groups:
                await m.answer("Пока нет прошлых командировок в истории. Создай новую командировку.", reply_markup=make_keyboard(["🆕 Новая командировка", "🏠 Старт"], cols=1))
                return
            temp_search[chat_id] = groups
            msg = "Выбери прошлую командировку для повтора:\n\n"
            for i, g in enumerate(groups, 1):
                msg += f"{i}. {g.get('start')} — {g.get('end')}\n"
                for emp in g.get("employees", []):
                    msg += f"   - {emp.get('fio')}\n"
                msg += "\n"
            state[chat_id] = "trip_repeat_choose"
            await m.answer(msg, reply_markup=trip_repeat_keyboard(groups))
            return
        await m.answer("Выбери действие кнопкой.")
        return

    if state.get(chat_id) == "trip_repeat_choose":
        if not text.isdigit():
            await m.answer("Выбери номер прошлой командировки кнопкой.")
            return
        groups = temp_search.get(chat_id, [])
        idx = int(text) - 1
        if idx < 0 or idx >= len(groups):
            await m.answer("Неверный номер. Выбери кнопкой.")
            return
        selected = groups[idx]
        employees = []
        for emp in selected.get("employees", []):
            pos = emp.get("position", "").replace(TRIP_REGION_SUFFIX, "").replace("Инженер-программист", "Инженер программист").strip()
            if pos not in ["Инженер программист", "Программист"]:
                pos = "Инженер программист"
            employees.append({"fio": emp.get("fio", ""), "pos": pos})
        data[chat_id] = {"employees": employees}
        state[chat_id] = "trip_repeat_edit"
        msg = "Выбрана командировка как шаблон. Сейчас список сотрудников такой:\n\n"
        for i, emp in enumerate(employees, 1):
            msg += f"{i}. {emp.get('fio')} — {trip_position_text(emp.get('pos', ''))}\n"
        await m.answer(msg + "\nМожешь добавить, удалить сотрудника или продолжить.", reply_markup=make_keyboard(["➕ Добавить сотрудника", "➖ Удалить сотрудника", "✅ Продолжить", "🏠 Старт"], cols=1))
        return

    if state.get(chat_id) == "trip_repeat_edit":
        if text == "➕ Добавить сотрудника":
            state[chat_id] = "trip_repeat_search_emp"
            await m.answer("Напиши минимум 2 буквы фамилии или имени.", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            return
        if text == "➖ Удалить сотрудника":
            employees = data[chat_id].get("employees", [])
            if len(employees) <= 1:
                await m.answer("Нельзя удалить: должен остаться хотя бы 1 сотрудник.")
                return
            msg = "Кого удалить из командировки?\n\n"
            buttons = []
            for i, emp in enumerate(employees, 1):
                msg += f"{i}. {emp.get('fio')}\n"
                buttons.append(str(i))
            buttons.append("🏠 Старт")
            state[chat_id] = "trip_repeat_delete_emp"
            await m.answer(msg, reply_markup=make_keyboard(buttons, cols=3))
            return
        if text == "✅ Продолжить":
            state[chat_id] = "trip_start_date"
            await m.answer("Введи новую дату начала командировки ДД.ММ.ГГГГ", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            return
        await m.answer("Выбери действие кнопкой.")
        return

    if state.get(chat_id) == "trip_repeat_delete_emp":
        if not text.isdigit():
            await m.answer("Выбери номер сотрудника кнопкой.")
            return
        employees = data[chat_id].get("employees", [])
        idx = int(text) - 1
        if idx < 0 or idx >= len(employees):
            await m.answer("Неверный номер. Выбери кнопкой.")
            return
        removed = employees.pop(idx)
        state[chat_id] = "trip_repeat_edit"
        await m.answer(f"Удален: {removed.get('fio')}", reply_markup=make_keyboard(["➕ Добавить сотрудника", "➖ Удалить сотрудника", "✅ Продолжить", "🏠 Старт"], cols=1))
        return

    if state.get(chat_id) in ["trip_repeat_search_emp", "trip_search_emp"]:
        found = search_employees_by_text(text, load_employees())
        if found is None:
            await m.answer("Напиши минимум 2 буквы фамилии или имени.")
            return
        already = {emp.get("fio") for emp in data[chat_id].get("employees", [])}
        found = [e for e in found if e not in already]
        if not found:
            await m.answer("Сотрудник не найден или уже выбран. Напиши другую часть ФИО.")
            return
        temp_search[chat_id] = found
        state[chat_id] = "trip_repeat_choose_emp" if state.get(chat_id) == "trip_repeat_search_emp" else "trip_choose_emp"
        await m.answer("Выбери сотрудника:", reply_markup=employee_keyboard(found))
        return

    if state.get(chat_id) in ["trip_repeat_choose_emp", "trip_choose_emp"]:
        if text not in temp_search.get(chat_id, []):
            await m.answer("Выбери сотрудника только из списка кнопок.")
            return
        data[chat_id]["selected_trip_fio"] = text
        state[chat_id] = "trip_repeat_pos" if state.get(chat_id) == "trip_repeat_choose_emp" else "trip_pos"
        await m.answer("Должность:", reply_markup=pos_menu)
        return

    if state.get(chat_id) in ["trip_repeat_pos", "trip_pos"]:
        if text not in ["Инженер программист", "Программист"]:
            await m.answer("Выбери должность из списка.")
            return
        data[chat_id]["employees"].append({"fio": data[chat_id].get("selected_trip_fio", ""), "pos": text})
        if state.get(chat_id) == "trip_repeat_pos":
            state[chat_id] = "trip_repeat_edit"
            await m.answer("Сотрудник добавлен ✅", reply_markup=make_keyboard(["➕ Добавить сотрудника", "➖ Удалить сотрудника", "✅ Продолжить", "🏠 Старт"], cols=1))
            return
        current = len(data[chat_id]["employees"])
        total = data[chat_id].get("trip_count", 1)
        if current < total:
            state[chat_id] = "trip_search_emp"
            await m.answer(f"Сотрудник {current + 1} из {total}. Напиши минимум 2 буквы фамилии или имени.", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            return
        state[chat_id] = "trip_start_date"
        await m.answer("Введи дату начала командировки ДД.ММ.ГГГГ", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
        return

    if state.get(chat_id) == "trip_count":
        if not text.isdigit():
            await m.answer("Напиши количество сотрудников цифрой.")
            return
        count = int(text)
        if count < 1 or count > 20:
            await m.answer("Количество должно быть от 1 до 20.")
            return
        data[chat_id] = {"employees": [], "trip_count": count}
        state[chat_id] = "trip_search_emp"
        await m.answer(f"Сотрудник 1 из {count}. Напиши минимум 2 буквы фамилии или имени.")
        return

    if state.get(chat_id) == "trip_start_date":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату начала в формате ДД.ММ.ГГГГ")
            return
        data[chat_id]["start"] = normalize_date(text)
        state[chat_id] = "trip_end_date"
        await m.answer("Введи дату конца командировки ДД.ММ.ГГГГ")
        return

    if state.get(chat_id) == "trip_end_date":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату конца в формате ДД.ММ.ГГГГ")
            return
        start_date = datetime.strptime(data[chat_id]["start"], "%d.%m.%Y")
        end_date = datetime.strptime(normalize_date(text), "%d.%m.%Y")
        if end_date < start_date:
            await m.answer("Дата конца не может быть раньше даты начала. Введи дату конца заново.")
            return
        data[chat_id]["end"] = normalize_date(text)
        if need_overlap_check(chat_id):
            overlaps = find_trip_overlaps(data[chat_id])
            if overlaps:
                data[chat_id]["pending_action"] = "trip"
                state[chat_id] = "overlap_confirm"
                await m.answer(format_trip_overlap_warning(overlaps, data[chat_id].get("start"), data[chat_id].get("end")), reply_markup=overlap_menu)
                return
        await finalize_trip_action(m, chat_id)
        return

    # ================== РУЧНАЯ ЗАПИСЬ ==================
    if text == "➕ Добавить запись вручную":
        state[chat_id] = "manual_search_emp"
        await m.answer("Напиши ФИО или часть ФИО сотрудника")
        return

    if state.get(chat_id) == "manual_search_emp":
        found = search_employees_by_text(text, load_employees())
        if found is None:
            await m.answer("Напиши минимум 2 буквы фамилии или имени.")
            return
        if not found:
            await m.answer("Сотрудник не найден. Напиши другую часть ФИО.")
            return
        temp_search[chat_id] = found
        state[chat_id] = "manual_choose_emp"
        await m.answer("Выбери сотрудника:", reply_markup=employee_keyboard(found))
        return

    if state.get(chat_id) == "manual_choose_emp":
        if text not in temp_search.get(chat_id, []):
            await m.answer("Выбери сотрудника только из списка кнопок.")
            return
        data[chat_id] = {"fio": text, "pos": ""}
        state[chat_id] = "manual_type"
        await m.answer("Выбери тип записи:", reply_markup=manual_type_menu)
        return

    if state.get(chat_id) == "manual_type":
        if text not in MANUAL_TYPES:
            await m.answer("Выбери тип записи из списка.")
            return
        data[chat_id]["type"] = text
        profile = get_employee_last_profile(data[chat_id].get("fio"))
        if profile:
            data[chat_id]["saved_profile"] = profile
            state[chat_id] = "manual_profile_confirm"
            await m.answer(profile_message(profile), reply_markup=saved_profile_menu)
            return
        state[chat_id] = "manual_pos"
        await m.answer("Должность:", reply_markup=pos_menu)
        return

    if state.get(chat_id) == "manual_profile_confirm":
        if text == "✅ Да, использовать":
            profile = data[chat_id].get("saved_profile", {})
            data[chat_id]["pos"] = profile.get("pos", "")
            data[chat_id]["project"] = profile.get("project", "")
            if not data[chat_id].get("pos"):
                state[chat_id] = "manual_pos"
                await m.answer("Должность:", reply_markup=pos_menu)
                return
            if not data[chat_id].get("project"):
                state[chat_id] = "manual_project"
                await m.answer("Напиши название проекта")
                return
            state[chat_id] = "manual_start"
            await m.answer("Введи дату начала ДД.ММ.ГГГГ", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            return
        if text == "✏️ Изменить":
            state[chat_id] = "manual_pos"
            await m.answer("Должность:", reply_markup=pos_menu)
            return
        await m.answer("Выбери действие кнопкой.", reply_markup=saved_profile_menu)
        return

    if state.get(chat_id) == "manual_pos":
        if text not in ["Инженер программист", "Программист"]:
            await m.answer("Выбери должность из списка.")
            return
        data[chat_id]["pos"] = text
        state[chat_id] = "manual_project"
        await m.answer("Напиши название проекта")
        return

    if state.get(chat_id) == "manual_project":
        data[chat_id]["project"] = text
        state[chat_id] = "manual_start"
        await m.answer("Введи дату начала ДД.ММ.ГГГГ")
        return

    if state.get(chat_id) == "manual_start":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату в формате ДД.ММ.ГГГГ")
            return
        data[chat_id]["start"] = normalize_date(text)
        if data[chat_id]["type"] == "📅 БС на один день":
            data[chat_id]["end"] = normalize_date(text)
            data[chat_id]["days"] = "1"
            if await check_overlap_or_continue(m, chat_id, "manual"):
                await finalize_manual_action(m, chat_id)
            return
        state[chat_id] = "manual_end"
        await m.answer("Введи дату конца ДД.ММ.ГГГГ")
        return

    if state.get(chat_id) == "manual_end":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату конца в формате ДД.ММ.ГГГГ")
            return
        start_date = datetime.strptime(data[chat_id]["start"], "%d.%m.%Y")
        end_date = datetime.strptime(normalize_date(text), "%d.%m.%Y")
        if end_date < start_date:
            await m.answer("Дата конца не может быть раньше даты начала. Введи дату конца заново.")
            return
        data[chat_id]["end"] = normalize_date(text)
        data[chat_id]["days"] = str((end_date - start_date).days + 1)
        if await check_overlap_or_continue(m, chat_id, "manual"):
            await finalize_manual_action(m, chat_id)
        return

    # ================== УДАЛЕНИЕ ==================
    if text == "🗑 Удалить запись":
        state[chat_id] = "delete_search_emp"
        await m.answer("Напиши ФИО или часть ФИО сотрудника")
        return

    if state.get(chat_id) == "delete_search_emp":
        found = search_fios_by_text(text, get_employees_and_history_fios())
        if found is None:
            await m.answer("Напиши минимум 2 буквы фамилии или имени.")
            return
        if not found:
            await m.answer("Сотрудник не найден. Напиши другую часть ФИО.")
            return
        temp_search[chat_id] = found
        state[chat_id] = "delete_choose_emp"
        await m.answer("Выбери сотрудника:", reply_markup=employee_keyboard(found))
        return

    if state.get(chat_id) == "delete_choose_emp":
        if text not in temp_search.get(chat_id, []):
            await m.answer("Выбери сотрудника только из списка кнопок.")
            return
        history = load_history()
        records = [r for r in history if r.get("fio") == text]
        if not records:
            records = [r for r in history if isinstance(r, dict) and fio_startswith_text(r.get("fio", ""), text)]
        if not records:
            state[chat_id] = "menu"
            await m.answer("У этого сотрудника нет записей в истории.", reply_markup=menu)
            return
        data[chat_id] = {"delete_fio": text, "delete_records": records}
        buttons = []
        msg = f"🗑 Записи сотрудника:\n\n👤 {text}\n\n"
        for i, r in enumerate(records, 1):
            msg += f"{i}. {r.get('type')}\n   Проект: {r.get('project', '')}\n"
            if r.get("periods"):
                msg += "   Периоды:\n"
                for p in get_periods_from_record(r):
                    msg += f"   - {p.get('start')} по {p.get('end')}\n"
            else:
                msg += f"   С: {r.get('start')} по {r.get('end')}\n"
            if r.get("days"):
                msg += f"   Дней: {r.get('days')}\n"
            if r.get("return_date"):
                msg += f"   Выход: {r.get('return_date')}\n"
            msg += f"   Создан: {r.get('created_at')}\n\n"
            buttons.append(str(i))
        buttons.append("🏠 Старт")
        state[chat_id] = "delete_choose_record"
        await m.answer(msg + "Выбери номер записи для удаления:", reply_markup=make_keyboard(buttons, cols=3))
        return

    if state.get(chat_id) == "delete_choose_record":
        if not text.isdigit():
            await m.answer("Выбери номер записи кнопкой.")
            return
        idx = int(text) - 1
        records = data[chat_id].get("delete_records", [])
        if idx < 0 or idx >= len(records):
            await m.answer("Неверный номер записи.")
            return
        record = records[idx]
        data[chat_id]["delete_record"] = record
        msg = (
            f"Ты точно хочешь удалить эту запись?\n\n"
            f"👤 {record.get('fio')}\n"
            f"Проект: {record.get('project', '')}\n"
            f"Тип: {record.get('type')}\n"
            f"С: {record.get('start')} по {record.get('end')}\n"
            f"Дней: {record.get('days')}\n"
            f"Выход: {record.get('return_date')}\n"
        )
        state[chat_id] = "delete_confirm"
        await m.answer(msg, reply_markup=confirm_delete_menu)
        return

    if state.get(chat_id) == "delete_confirm":
        if text == "❌ Нет, отменить":
            state[chat_id] = "menu"
            await m.answer("Удаление отменено.", reply_markup=menu)
            return
        if text != "✅ Да, удалить":
            await m.answer("Выбери: ✅ Да, удалить или ❌ Нет, отменить")
            return
        record_to_delete = data[chat_id].get("delete_record")
        history = load_history()
        new_history = []
        deleted = False
        for r in history:
            if not deleted and r == record_to_delete:
                deleted = True
                continue
            new_history.append(r)
        save_history_full(new_history)
        state[chat_id] = "menu"
        await m.answer("Запись удалена ✅", reply_markup=menu)
        return

    # ================== БОЛЬНИЧНЫЙ ==================
    if text == "🏥 Больничный":
        state[chat_id] = "sick_search_emp"
        await m.answer("Напиши ФИО или часть ФИО сотрудника")
        return

    if state.get(chat_id) == "sick_search_emp":
        if len(str(text).lower().strip()) < 2:
            await m.answer("Напиши минимум 2 буквы фамилии или имени.")
            return
        active = find_active_sick_records_by_text(text)
        if len(active) == 1:
            old_record = active[0]
            data[chat_id] = {"fio": old_record.get("fio"), "type": SICK_LEAVE_TYPE, "old_sick_record": old_record}
            state[chat_id] = "sick_extend_start_date"
            await m.answer(f"✅ Найдена активная запись больничного:\n\n👤 {old_record.get('fio')}\nПроект: {old_record.get('project', '')}\nС: {old_record.get('start')} по {old_record.get('end')}\nДней: {old_record.get('days')}\nВыход: {old_record.get('return_date')}\n\nТеперь введи дату, С КОТОРОЙ продлеваем больничный: ДД.ММ.ГГГГ")
            return
        if len(active) > 1:
            found = unique_list([r.get("fio") for r in active])
            temp_search[chat_id] = found
            state[chat_id] = "sick_choose_emp"
            await m.answer("Найдено несколько активных больничных. Выбери сотрудника:", reply_markup=employee_keyboard(found))
            return
        found = search_employees_by_text(text, load_employees())
        if found is None:
            await m.answer("Напиши минимум 2 буквы фамилии или имени.")
            return
        if not found:
            await m.answer("Сотрудник не найден. Напиши другую часть ФИО.")
            return
        if len(found) == 1:
            data[chat_id] = {"fio": found[0], "type": SICK_LEAVE_TYPE, "pos": clean_position_for_profile(get_position_from_salary(found[0]))}
            state[chat_id] = "sick_project"
            await m.answer(f"Активный больничный по сотруднику {found[0]} не найден.\nСоздаем новую запись больничного. Напиши название проекта")
            return
        temp_search[chat_id] = found
        state[chat_id] = "sick_choose_emp"
        await m.answer("Активный больничный не найден. Выбери сотрудника для новой записи:", reply_markup=employee_keyboard(found))
        return

    if state.get(chat_id) == "sick_choose_emp":
        if text not in temp_search.get(chat_id, []):
            await m.answer("Выбери сотрудника только из списка кнопок.")
            return
        active = find_active_sick_records_by_fio(text)
        if active:
            old_record = active[0]
            data[chat_id] = {"fio": text, "type": SICK_LEAVE_TYPE, "old_sick_record": old_record}
            state[chat_id] = "sick_extend_start_date"
            await m.answer(f"✅ Найдена активная запись больничного:\n\n👤 {old_record.get('fio')}\nПроект: {old_record.get('project', '')}\nС: {old_record.get('start')} по {old_record.get('end')}\nДней: {old_record.get('days')}\nВыход: {old_record.get('return_date')}\n\nТеперь введи дату, С КОТОРОЙ продлеваем больничный: ДД.ММ.ГГГГ")
            return
        data[chat_id] = {"fio": text, "type": SICK_LEAVE_TYPE, "pos": clean_position_for_profile(get_position_from_salary(text))}
        state[chat_id] = "sick_project"
        await m.answer("Активный больничный не найден. Создаем новую запись. Напиши название проекта")
        return

    if state.get(chat_id) == "sick_extend_start_date":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату начала продления в формате ДД.ММ.ГГГГ")
            return
        old_record = data[chat_id].get("old_sick_record")
        old_end = datetime.strptime(old_record.get("end"), "%d.%m.%Y")
        new_start = datetime.strptime(normalize_date(text), "%d.%m.%Y")
        if new_start <= old_end:
            await m.answer(f"Дата продления должна быть после текущей даты окончания больничного: {old_record.get('end')}.\nВведи дату начала продления заново.")
            return
        data[chat_id]["extend_start"] = normalize_date(text)
        state[chat_id] = "sick_extend_end_date"
        await m.answer("Введи дату, ДО КОТОРОЙ продлеваем больничный: ДД.ММ.ГГГГ")
        return

    if state.get(chat_id) == "sick_extend_end_date":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату конца продления в формате ДД.ММ.ГГГГ")
            return
        start_date = datetime.strptime(data[chat_id]["extend_start"], "%d.%m.%Y")
        end_date = datetime.strptime(normalize_date(text), "%d.%m.%Y")
        if end_date < start_date:
            await m.answer("Дата конца не может быть раньше даты начала. Введи дату конца заново.")
            return
        old_record = data[chat_id].get("old_sick_record")
        new_record = build_extended_record(old_record, data[chat_id]["extend_start"], normalize_date(text))
        update_history_record(old_record, new_record)
        await notify_sick_extended(old_record, new_record, data[chat_id]["extend_start"], normalize_date(text))
        state[chat_id] = "menu"
        await m.answer(f"Больничный продлен ✅\n\n👤 {new_record.get('fio')}\nВсего дней больничного: {new_record.get('days')}\nВыход на работу: {new_record.get('return_date')}", reply_markup=menu)
        return

    if state.get(chat_id) == "sick_project":
        data[chat_id]["project"] = text
        state[chat_id] = "sick_start_date"
        await m.answer("Введи дату начала больничного ДД.ММ.ГГГГ")
        return

    if state.get(chat_id) == "sick_start_date":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату в формате ДД.ММ.ГГГГ")
            return
        data[chat_id]["start"] = normalize_date(text)
        state[chat_id] = "sick_end_date"
        await m.answer("Введи примерную дату конца больничного ДД.ММ.ГГГГ")
        return

    if state.get(chat_id) == "sick_end_date":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату конца в формате ДД.ММ.ГГГГ")
            return
        start_date = datetime.strptime(data[chat_id]["start"], "%d.%m.%Y")
        end_date = datetime.strptime(normalize_date(text), "%d.%m.%Y")
        if end_date < start_date:
            await m.answer("Дата конца не может быть раньше даты начала. Введи дату конца заново.")
            return
        data[chat_id]["end"] = normalize_date(text)
        data[chat_id]["days"] = str((end_date - start_date).days + 1)
        data[chat_id]["periods"] = [{"start": data[chat_id]["start"], "end": normalize_date(text)}]
        if await check_overlap_or_continue(m, chat_id, "sick"):
            await finalize_sick_action(m, chat_id)
        return

    # ================== ПРОДЛЕНИЕ БС ==================
    if state.get(chat_id) == "bs_extend_start_date":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату начала продления БС в формате ДД.ММ.ГГГГ")
            return
        old_record = data[chat_id].get("old_bs_record")
        old_end = datetime.strptime(old_record.get("end"), "%d.%m.%Y")
        new_start = datetime.strptime(normalize_date(text), "%d.%m.%Y")
        if new_start <= old_end:
            await m.answer(f"Дата продления должна быть после текущей даты окончания БС: {old_record.get('end')}.\nВведи дату начала продления заново.")
            return
        data[chat_id]["extend_start"] = normalize_date(text)
        state[chat_id] = "bs_extend_end_date"
        await m.answer("Введи дату, ДО КОТОРОЙ продлеваем БС: ДД.ММ.ГГГГ")
        return

    if state.get(chat_id) == "bs_extend_end_date":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату конца продления БС в формате ДД.ММ.ГГГГ")
            return
        start_date = datetime.strptime(data[chat_id]["extend_start"], "%d.%m.%Y")
        end_date = datetime.strptime(normalize_date(text), "%d.%m.%Y")
        if end_date < start_date:
            await m.answer("Дата конца не может быть раньше даты начала. Введи дату конца заново.")
            return
        old_record = data[chat_id].get("old_bs_record")
        new_record = build_extended_record(old_record, data[chat_id]["extend_start"], normalize_date(text), force_type=BS_RANGE_TYPE)
        update_history_record(old_record, new_record)
        await notify_bs_extended(old_record, new_record, data[chat_id]["extend_start"], normalize_date(text))
        state[chat_id] = "menu"
        await m.answer(f"БС продлен ✅\n\n👤 {new_record.get('fio')}\nВсего дней БС: {new_record.get('days')}\nВыход на работу: {new_record.get('return_date')}", reply_markup=menu)
        return

    # ================== СОЗДАТЬ ЗАЯВЛЕНИЕ ==================
    if text == "📄 Создать заявление":
        state[chat_id] = "search_emp"
        await m.answer("Напиши ФИО или часть ФИО сотрудника")
        return

    if state.get(chat_id) == "search_emp":
        found = search_employees_by_text(text, load_employees())
        if found is None:
            await m.answer("Напиши минимум 2 буквы фамилии или имени.")
            return
        if not found:
            await m.answer("Сотрудник не найден. Напиши другую часть ФИО.")
            return
        temp_search[chat_id] = found
        state[chat_id] = "choose_emp"
        await m.answer("Выбери сотрудника:", reply_markup=employee_keyboard(found))
        return

    if state.get(chat_id) == "choose_emp":
        if text not in temp_search.get(chat_id, []):
            await m.answer("Выбери сотрудника только из списка кнопок.")
            return
        data[chat_id] = {"fio": text}
        state[chat_id] = "type"
        await m.answer("Выбери тип заявления:", reply_markup=make_keyboard(list(TEMPLATES.keys()) + [SICK_LEAVE_TYPE, "🏠 Старт"], cols=2))
        return

    if state.get(chat_id) == "type":
        if text not in TEMPLATES and text != SICK_LEAVE_TYPE:
            await m.answer("Выбери тип заявления из списка.")
            return
        data[chat_id]["type"] = text

        if text == SICK_LEAVE_TYPE:
            active = find_active_sick_records_by_fio(data[chat_id].get("fio"))
            if active:
                old_record = active[0]
                data[chat_id] = {"fio": old_record.get("fio"), "type": SICK_LEAVE_TYPE, "old_sick_record": old_record}
                state[chat_id] = "sick_extend_start_date"
                await m.answer(f"✅ Найдена активная запись больничного:\n\n👤 {old_record.get('fio')}\nПроект: {old_record.get('project', '')}\nС: {old_record.get('start')} по {old_record.get('end')}\nДней: {old_record.get('days')}\nВыход: {old_record.get('return_date')}\n\nТеперь введи дату, С КОТОРОЙ продлеваем больничный: ДД.ММ.ГГГГ")
                return
            data[chat_id] = {
                "fio": data[chat_id].get("fio"),
                "type": SICK_LEAVE_TYPE,
                "pos": clean_position_for_profile(get_position_from_salary(data[chat_id].get("fio"))),
            }
            state[chat_id] = "sick_project"
            await m.answer("Активный больничный не найден. Создаем новую запись больничного. Напиши название проекта")
            return

        if text == "🧩 Часть отпуска":
            previous = find_previous_part_leave_records_by_fio(data[chat_id].get("fio"))
            if previous:
                data[chat_id]["type"] = "📌 Оставшийся отпуск"
                await m.answer("ℹ️ У этого сотрудника уже есть запись по части отпуска.\n\n" + f"👤 {data[chat_id].get('fio')}\n" + format_leave_records_for_message(previous) + "\n✅ Тип заявления автоматически изменен на: 📌 Оставшийся отпуск")
        if text in BS_LEAVE_TYPES:
            active = find_active_bs_records_by_fio(data[chat_id].get("fio"))
            if active:
                old_record = active[0]
                data[chat_id] = {"fio": old_record.get("fio"), "type": BS_RANGE_TYPE, "old_bs_record": old_record}
                state[chat_id] = "bs_extend_start_date"
                await m.answer(f"✅ Найдена активная запись БС:\n\n👤 {old_record.get('fio')}\nПроект: {old_record.get('project', '')}\nС: {old_record.get('start')} по {old_record.get('end')}\nДней: {old_record.get('days')}\nВыход: {old_record.get('return_date')}\n\nТеперь введи дату, С КОТОРОЙ продлеваем БС: ДД.ММ.ГГГГ")
                return
        profile = get_employee_last_profile(data[chat_id].get("fio"))
        if profile:
            data[chat_id]["saved_profile"] = profile
            state[chat_id] = "profile_confirm"
            await m.answer(profile_message(profile), reply_markup=saved_profile_menu)
            return
        state[chat_id] = "pos"
        await m.answer("Должность:", reply_markup=pos_menu)
        return

    if state.get(chat_id) == "profile_confirm":
        if text == "✅ Да, использовать":
            profile = data[chat_id].get("saved_profile", {})
            data[chat_id]["pos"] = profile.get("pos", "")
            data[chat_id]["project"] = profile.get("project", "")
            if not data[chat_id].get("pos"):
                state[chat_id] = "pos"
                await m.answer("Должность:", reply_markup=pos_menu)
                return
            if not data[chat_id].get("project"):
                state[chat_id] = "project"
                await m.answer("Напиши название проекта")
                return
            if data[chat_id]["type"] in MATERIAL_ASSISTANCE_NO_DATE_TYPES:
                if await check_overlap_or_continue(m, chat_id, "doc"):
                    await finalize_doc_action(m, chat_id)
                return
            state[chat_id] = "date"
            await m.answer("Введи дату начала ДД.ММ.ГГГГ", reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            return
        if text == "✏️ Изменить":
            state[chat_id] = "pos"
            await m.answer("Должность:", reply_markup=pos_menu)
            return
        await m.answer("Выбери действие кнопкой.", reply_markup=saved_profile_menu)
        return

    if state.get(chat_id) == "pos":
        if text not in ["Инженер программист", "Программист"]:
            await m.answer("Выбери должность из списка.")
            return
        data[chat_id]["pos"] = text
        state[chat_id] = "project"
        await m.answer("Напиши название проекта")
        return

    if state.get(chat_id) == "project":
        data[chat_id]["project"] = text
        if data[chat_id]["type"] in MATERIAL_ASSISTANCE_NO_DATE_TYPES:
            if await check_overlap_or_continue(m, chat_id, "doc"):
                await finalize_doc_action(m, chat_id)
            return
        state[chat_id] = "date"
        await m.answer("Введи дату начала ДД.ММ.ГГГГ")
        return

    if state.get(chat_id) == "date":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату в формате ДД.ММ.ГГГГ")
            return
        data[chat_id]["start"] = normalize_date(text)
        start_date = datetime.strptime(normalize_date(text), "%d.%m.%Y")
        doc_type = data[chat_id]["type"]
        if doc_type == "🌴 Полный отпуск":
            days = 30
            data[chat_id]["days"] = str(days)
            data[chat_id]["end"] = (start_date + timedelta(days=days - 1)).strftime("%d.%m.%Y")
            if await check_overlap_or_continue(m, chat_id, "doc"):
                await finalize_doc_action(m, chat_id)
            return
        if doc_type == "🧩 Часть отпуска":
            days = 15
            data[chat_id]["days"] = str(days)
            data[chat_id]["end"] = (start_date + timedelta(days=days - 1)).strftime("%d.%m.%Y")
            if await check_overlap_or_continue(m, chat_id, "doc"):
                await finalize_doc_action(m, chat_id)
            return
        if doc_type == "📚 Учебный отпуск":
            state[chat_id] = "study_end_date"
            await m.answer("Введи дату конца учебного отпуска ДД.ММ.ГГГГ")
            return
        if doc_type == "📅 БС на один день":
            data[chat_id]["days"] = "1"
            data[chat_id]["end"] = normalize_date(text)
            if await check_overlap_or_continue(m, chat_id, "doc"):
                await finalize_doc_action(m, chat_id)
            return
        if doc_type in ["💍 Мат помощь (свадьба)", CHILD_BIRTH_3_DAYS_TYPE]:
            days = 3
            data[chat_id]["days"] = str(days)
            data[chat_id]["end"] = (start_date + timedelta(days=days - 1)).strftime("%d.%m.%Y")
            if await check_overlap_or_continue(m, chat_id, "doc"):
                await finalize_doc_action(m, chat_id)
            return
        state[chat_id] = "days"
        await m.answer("Количество дней")
        return

    if state.get(chat_id) == "study_end_date":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату конца в формате ДД.ММ.ГГГГ")
            return
        start_date = datetime.strptime(data[chat_id]["start"], "%d.%m.%Y")
        end_date = datetime.strptime(normalize_date(text), "%d.%m.%Y")
        if end_date < start_date:
            await m.answer("Дата конца не может быть раньше даты начала. Введи дату конца заново.")
            return
        data[chat_id]["end"] = normalize_date(text)
        data[chat_id]["days"] = str((end_date - start_date).days + 1)
        if await check_overlap_or_continue(m, chat_id, "doc"):
            await finalize_doc_action(m, chat_id)
        return

    if state.get(chat_id) == "days":
        if not text.isdigit():
            await m.answer("Введи количество дней цифрой.")
            return
        days = int(text)
        start_date = datetime.strptime(data[chat_id]["start"], "%d.%m.%Y")
        data[chat_id]["days"] = str(days)
        data[chat_id]["end"] = (start_date + timedelta(days=days - 1)).strftime("%d.%m.%Y")
        if await check_overlap_or_continue(m, chat_id, "doc"):
            await finalize_doc_action(m, chat_id)
        return

    # ================== ИСТОРИЯ ==================
    if text == "📜 История":
        state[chat_id] = "history_menu"
        await m.answer("Выбери действие:", reply_markup=history_menu)
        return

    if text == "🔍 Поиск сотрудника":
        state[chat_id] = "history_search"
        await m.answer("Напиши ФИО или часть ФИО сотрудника")
        return

    if text == "📋 Полный список сотрудников":
        fios = get_employees_and_history_fios()
        temp_search[chat_id] = fios
        state[chat_id] = "history_choose"
        await m.answer("Выбери сотрудника:", reply_markup=employee_keyboard(fios))
        return

    if state.get(chat_id) == "history_search":
        found = search_fios_by_text(text, get_employees_and_history_fios())
        if found is None:
            await m.answer("Напиши минимум 2 буквы фамилии или имени.")
            return
        if not found:
            await m.answer("Сотрудник не найден. Попробуй ещё раз.")
            return
        temp_search[chat_id] = found
        state[chat_id] = "history_choose"
        await m.answer("Выбери сотрудника:", reply_markup=employee_keyboard(found))
        return

    if state.get(chat_id) == "history_choose":
        if text not in temp_search.get(chat_id, []):
            await m.answer("Выбери сотрудника только из списка кнопок.")
            return
        history = load_history()
        records = [r for r in history if r.get("fio") == text]
        if not records:
            records = [r for r in history if isinstance(r, dict) and fio_startswith_text(r.get("fio", ""), text)]
        if not records:
            await m.answer("По этому сотруднику история пустая.", reply_markup=menu)
            state[chat_id] = "menu"
            return
        msg = ""
        for r in records:
            msg += f"ФИО: {r.get('fio')}\nТип: {r.get('type')}\nДолжность: {display_position(r.get('fio', ''), r.get('position', ''))}\n"
            if r.get("type") != TRIP_TYPE:
                msg += f"Проект: {r.get('project', '')}\n"
            if r.get("periods"):
                msg += "Периоды:\n"
                for p in get_periods_from_record(r):
                    msg += f"- {p.get('start')} по {p.get('end')}\n"
            else:
                if r.get("start"):
                    msg += f"Начало: {r.get('start')}\n"
                if r.get("end"):
                    msg += f"Конец: {r.get('end')}\n"
            if r.get("days"):
                msg += f"Дней: {r.get('days')}\n"
            if r.get("return_date"):
                msg += f"Выход на работу: {r.get('return_date')}\n"
            msg += f"Создан: {r.get('created_at')}\n\n"

        changeable_records = [r for r in records if can_change_return_date(r)]
        data[chat_id] = {"history_fio": text, "history_records": records, "return_change_records": changeable_records}
        state[chat_id] = "history_action"
        if changeable_records:
            await m.answer(msg, reply_markup=history_action_menu)
        else:
            await m.answer(msg, reply_markup=menu)
            state[chat_id] = "menu"
        return

    if state.get(chat_id) == "history_action":
        if text == "✏️ Изменить дату выхода":
            records = data.get(chat_id, {}).get("return_change_records", [])
            record = choose_return_change_record(records)
            if not record:
                state[chat_id] = "menu"
                await m.answer("По этому сотруднику нет записи с датой выхода.", reply_markup=menu)
                return
            data[chat_id]["return_change_record"] = record
            state[chat_id] = "return_change_new_date"
            await m.answer(return_change_prompt(record), reply_markup=make_keyboard(["🏠 Старт"], cols=1))
            return
        await m.answer("Выбери действие кнопкой.", reply_markup=history_action_menu)
        return

    if state.get(chat_id) == "return_change_new_date":
        if not is_valid_date(text):
            await m.answer("Ошибка. Введи дату выхода в формате ДД.ММ.ГГГГ")
            return
        old_record = data.get(chat_id, {}).get("return_change_record")
        if not old_record:
            state[chat_id] = "menu"
            await m.answer("Запись не найдена. Начни заново через историю.", reply_markup=menu)
            return
        old_return = normalize_date(old_record.get("return_date", ""))
        new_return = normalize_date(text)
        if old_return == new_return:
            state[chat_id] = "menu"
            await m.answer("Дата выхода не изменилась.", reply_markup=menu)
            return
        new_record = dict(old_record)
        new_record["return_date"] = new_return
        new_record["return_changed_at"] = now_dt().strftime("%d.%m.%Y %H:%M:%S")
        new_record["return_change_reason"] = "дополнительный выходной / праздничный день"
        update_history_record(old_record, new_record)
        await notify_return_date_changed(old_record, new_record)
        state[chat_id] = "menu"
        await m.answer(
            f"Дата выхода изменена ✅\n\n👤 {new_record.get('fio')}\nБыло: {old_return}\nСтало: {new_return}",
            reply_markup=menu
        )
        return

    await m.answer("Выбери действие из меню.", reply_markup=get_menu(chat_id))


WEBHOOK_PATH = "/webhook"
WEBHOOK_URL = os.environ.get("WEBHOOK_URL", "https://office-bot-production-c91b.up.railway.app")
WEBHOOK_SECRET = os.environ.get("WEBHOOK_SECRET", "office_bot_secret")


async def webhook_handler(request):
    try:
        if request.headers.get("X-Telegram-Bot-Api-Secret-Token") != WEBHOOK_SECRET:
            return web.Response(status=403, text="Forbidden")
        update_data = await request.json()
        update = Update.model_validate(update_data, context={"bot": bot})
        await dp.feed_update(bot, update)
        return web.Response(text="OK")
    except Exception as e:
        print("WEBHOOK ERROR:", e)
        return web.Response(text="OK")


async def home(request):
    return web.Response(text="OK")


async def on_startup(app):
    print("BOT STARTING...")
    init_database()
    migrate_json_history_to_postgres()
    migrate_json_reminders_to_postgres()
    normalize_history_file()
    load_salary_records()
    load_sent_reminders()
    print("CACHE READY: history, salary, reminders загружены")
    webhook_full_url = WEBHOOK_URL.rstrip("/") + WEBHOOK_PATH
    await bot.delete_webhook(drop_pending_updates=True)
    await bot.set_webhook(webhook_full_url, secret_token=WEBHOOK_SECRET, drop_pending_updates=True)
    app["reminder_task"] = asyncio.create_task(reminder_loop())
    print("WEBHOOK SET:", webhook_full_url)
    print("BOT STARTED")


async def on_shutdown(app):
    print("BOT SHUTTING DOWN...")
    task = app.get("reminder_task")
    if task:
        task.cancel()
        try:
            await task
        except asyncio.CancelledError:
            pass
    await bot.session.close()
    print("BOT STOPPED")


def main():
    port = int(os.environ.get("PORT", 10000))
    app = web.Application()
    app.router.add_get("/", home)
    app.router.add_post(WEBHOOK_PATH, webhook_handler)
    app.on_startup.append(on_startup)
    app.on_shutdown.append(on_shutdown)
    print("WEB SERVER STARTING ON PORT", port)
    web.run_app(app, host="0.0.0.0", port=port)


if __name__ == "__main__":
    main()
